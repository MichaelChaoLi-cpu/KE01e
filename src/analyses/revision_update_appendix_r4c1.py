#!/usr/bin/env python3
"""Update approved Appendix A1/B1 content for Reviewer 4 Comment 1."""

from __future__ import annotations

from copy import deepcopy
from datetime import datetime, timezone
import hashlib
import os
from pathlib import Path
import shutil
import tempfile
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
import pandas as pd


ROOT = Path(__file__).resolve().parents[2]
APPENDIX = ROOT / "Rev/revision/Appendix.docx"
BACKUP_DIR = ROOT / "Rev/revision/.kila-backups"
COVERAGE = ROOT / "data/results/tables/Table_analytical_data_coverage_and_quality.xlsx"
HAZARD = ROOT / "data/results/tables/Table_hazard_validation.xlsx"
OLD_B1_PARAGRAPH = (
    "Appendix Table B1 reports the spatial validation results after excluding "
    "post-event and unknown-date warning zones from the 2016 validation. Across "
    "five folds, the fitted terrain-plus-warning model has a mean area under the "
    "curve of 0.599 and a fold range of 0.405–0.770, so it fails the stability "
    "rule. The fixed standardized terrain score has a mean of 0.665, a fold range "
    "of 0.581–0.745, and held-out top-quartile capture of 46.4%, so it is selected "
    "as the transparent scenario ranking. The warning-zone-only indicator has "
    "mean area under the curve of 0.556 and rank correlation 0.008 with the full "
    "ranking. All specifications use 857 presence cells, 8,570 reproducibly "
    "sampled pseudo-background cells, and 30,021 temporally eligible pre-event zones."
)
NEW_B1_PARAGRAPH = (
    "Appendix Table B1 reports the footprint-bounded historical terrain-ranking "
    "comparison after restricting warning-zone exposure to 29,632 polygons "
    "designated by 14 April 2016 and sampling pseudo-background only within the "
    "57.8% GSI interpretation footprint. Across five folds, the frozen fixed "
    "standardized terrain score has a mean spatial AUC of 0.705, a fold range of "
    "0.550–0.787, and held-out top-quartile capture of 46.5%. The fitted full "
    "terrain-plus-warning comparator has a mean AUC of 0.685 and a fold range of "
    "0.500–0.804, whereas the warning-zone-only indicator has a mean AUC of 0.513 "
    "and rank correlation of 0.020 with the full ranking. All specifications use "
    "857 presence cells and 8,570 reproducibly sampled pseudo-background cells. "
    "The fixed score was evaluated without refitting, so the support correction "
    "does not change the downstream scenario results."
)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def package_members(path: Path) -> dict[str, bytes]:
    with ZipFile(path) as archive:
        damaged = archive.testzip()
        if damaged:
            raise RuntimeError(f"Corrupt DOCX member: {damaged}")
        return {name: archive.read(name) for name in archive.namelist()}


def display_value(value: object, column: str) -> str:
    if pd.isna(value):
        return ""
    if column == "Record Count":
        return f"{int(value):,}"
    if column in {"Required-Field Missingness", "Location Completeness"}:
        return f"{100 * float(value):.4f}%" if abs(float(value)) < 0.001 else f"{100 * float(value):.1f}%"
    if column == "Spatial Folds":
        return f"{int(value)}"
    if column in {"Mean Spatial AUC", "Held-Out Top-Quartile Capture"}:
        return f"{100 * float(value):.1f}%"
    if column == "Rank Correlation vs Full":
        return f"{float(value):.3f}"
    return str(value)


def set_cell_text(cell, text: str) -> None:
    """Replace visible cell text while preserving cell and paragraph properties."""
    paragraph = cell.paragraphs[0]
    template_rpr = None
    for run in paragraph.runs:
        if run._r.rPr is not None:
            template_rpr = deepcopy(run._r.rPr)
            break
    paragraph.clear()
    run = paragraph.add_run(text)
    if template_rpr is not None:
        if run._r.rPr is not None:
            run._r.remove(run._r.rPr)
        run._r.insert(0, template_rpr)
    for extra in list(cell.paragraphs[1:]):
        extra._element.getparent().remove(extra._element)


def update_table(table, frame: pd.DataFrame) -> None:
    expected = (len(frame) + 1, len(frame.columns))
    actual = (len(table.rows), len(table.columns))
    if actual != expected:
        raise RuntimeError(f"Appendix table shape {actual} does not match {expected}.")
    for column, label in enumerate(frame.columns):
        set_cell_text(table.cell(0, column), str(label))
    for row_index, values in enumerate(frame.itertuples(index=False, name=None), start=1):
        for column, (label, value) in enumerate(zip(frame.columns, values)):
            set_cell_text(table.cell(row_index, column), display_value(value, label))


def paragraph_visible_text(paragraph) -> str:
    """Return Word and native-math visible text in document order."""
    visible_tags = {
        qn("w:t"),
        qn("w:delText"),
        qn("m:t"),
    }
    return "".join(
        element.text or ""
        for element in paragraph._p.iter()
        if element.tag in visible_tags
    )


def replace_b1_explanatory_paragraph(document: Document) -> None:
    matches = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph_visible_text(paragraph) == OLD_B1_PARAGRAPH
    ]
    if len(matches) != 1:
        raise RuntimeError(
            f"Expected one exact Appendix B1 explanatory paragraph; found {len(matches)}."
        )
    paragraph = matches[0]
    protected_tags = {qn("w:pPr"), qn("w:bookmarkStart"), qn("w:bookmarkEnd")}
    for child in list(paragraph._p):
        if child.tag not in protected_tags:
            paragraph._p.remove(child)

    bookmark_end = next(
        (child for child in paragraph._p if child.tag == qn("w:bookmarkEnd")), None
    )
    if bookmark_end is None:
        raise RuntimeError("Appendix B1 paragraph bookmark end was not preserved.")
    run = paragraph._p.makeelement(qn("w:r"))
    text = run.makeelement(qn("w:t"))
    text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    text.text = NEW_B1_PARAGRAPH
    run.append(text)
    bookmark_end.addprevious(run)

    if paragraph_visible_text(paragraph) != NEW_B1_PARAGRAPH:
        raise RuntimeError("Appendix B1 paragraph replacement verification failed.")
    bookmark_ids = [
        child.get(qn("w:id"))
        for child in paragraph._p
        if child.tag in {qn("w:bookmarkStart"), qn("w:bookmarkEnd")}
    ]
    if bookmark_ids != ["3", "3"]:
        raise RuntimeError(f"Appendix B1 paragraph bookmark changed: {bookmark_ids}")


def main() -> None:
    coverage = pd.read_excel(COVERAGE, sheet_name="Data Coverage", header=1)
    hazard = pd.read_excel(HAZARD, sheet_name="Hazard Validation", header=1)
    if coverage.shape != (22, 10) or hazard.shape != (5, 8):
        raise RuntimeError(
            f"Unexpected formal table shapes: coverage={coverage.shape}, hazard={hazard.shape}"
        )

    before_hash = sha256(APPENDIX)
    before_members = package_members(APPENDIX)
    document = Document(APPENDIX)
    if len(document.tables) < 2:
        raise RuntimeError("Appendix does not contain Tables A1 and B1.")
    if document.tables[0].cell(0, 0).text != "Analytical Data Layer":
        raise RuntimeError("Appendix Table A1 anchor did not match.")
    if document.tables[1].cell(0, 0).text != "Specification":
        raise RuntimeError("Appendix Table B1 anchor did not match.")

    update_table(document.tables[0], coverage)
    update_table(document.tables[1], hazard)
    replace_b1_explanatory_paragraph(document)

    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    backup = BACKUP_DIR / f"Appendix.before-r4c1.{timestamp}.docx"
    shutil.copy2(APPENDIX, backup)
    with tempfile.NamedTemporaryFile(
        dir=APPENDIX.parent, prefix=".Appendix.r4c1.", suffix=".docx", delete=False
    ) as handle:
        staged = Path(handle.name)
    try:
        document.save(staged)
        staged_members = package_members(staged)
        allowed_changes = {"word/document.xml", "docProps/core.xml"}
        changed_members = {
            name
            for name in set(before_members) | set(staged_members)
            if before_members.get(name) != staged_members.get(name)
        }
        unexpected = changed_members - allowed_changes
        if unexpected:
            raise RuntimeError(f"Unexpected DOCX package changes: {sorted(unexpected)}")

        verified = Document(staged)
        if "57.8% air-photo footprint" not in verified.tables[0].cell(10, 3).text:
            raise RuntimeError("Updated Appendix Table A1 support text was not preserved.")
        if verified.tables[1].cell(5, 3).text != "70.5%":
            raise RuntimeError("Updated Appendix Table B1 frozen-score AUC was not preserved.")
        if "29,632 pre-sequence zones" not in verified.tables[1].cell(1, 1).text:
            raise RuntimeError("Updated Appendix Table B1 sample support was not preserved.")
        verified_matches = [
            paragraph
            for paragraph in verified.paragraphs
            if paragraph_visible_text(paragraph) == NEW_B1_PARAGRAPH
        ]
        if len(verified_matches) != 1:
            raise RuntimeError("Updated Appendix B1 explanatory paragraph was not preserved.")
        os.replace(staged, APPENDIX)
    finally:
        if staged.exists():
            staged.unlink()

    print(f"Appendix before SHA-256: {before_hash}")
    print(f"Appendix after SHA-256:  {sha256(APPENDIX)}")
    print(f"Backup: {backup.relative_to(ROOT)}")
    print("Updated only approved Appendix A1/B1 content; package-member guard passed.")


if __name__ == "__main__":
    main()
