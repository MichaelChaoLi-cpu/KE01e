#!/usr/bin/env python3
"""Apply the approved Reviewer 4 Comment 4 Appendix update and Table B12."""

from __future__ import annotations

import argparse
import datetime as dt
import hashlib
import json
import os
from pathlib import Path
import shutil
import tempfile
import zipfile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from openpyxl import load_workbook

from revision_update_appendix_r1c2 import replace_visible_content, visible_text
from revision_update_appendix_r2c5 import build_table


OLD_PARAGRAPH = (
    "Simulation-size and network-target checks bound the disconnection results. "
    "Changing the number of draws from 500 to 2,000 with one common seed produces "
    "a 0.006 difference in the 95th-percentile community isolation frequency. The "
    "Primary Emergency Road backbone contains 2,562 target roots and yields 1,063.6 "
    "expected disconnected residents under Heavy rainfall. Expanding the target to "
    "2,977 Primary-plus-Secondary Emergency Road roots yields 992.7 residents, with "
    "community-frequency Spearman correlation 0.964 and top-30 burden overlap 90.0% "
    "relative to the primary definition. The former coast-inclusive boundary proxy "
    "is reported only as an audit comparator and reproduces the prior 1,121.7-resident "
    "result. Municipality-wide Yatsushiro assignments of 0.70 and 0.80 bound the revised "
    "primary result between 1,016.6 and 1,118.7 residents (−4.4% to +5.2% relative to the "
    "0.75 midpoint); community-frequency rank correlations remain 0.989–1.000 and Top-30 "
    "population-burden overlap is 93.3%. Alternative closure mappings produce the wider "
    "rounded range of 343–2,057 residents. Closure mapping is therefore a larger source "
    "of magnitude uncertainty than Monte Carlo convergence or the tested emergency-road "
    "target definition."
)
NEW_PARAGRAPH = (
    "Simulation-size and network-target checks bound the disconnection results. "
    "Changing the number of draws from 500 to 2,000 with one common seed produces "
    "a 0.006 difference in the 95th-percentile community isolation frequency. The "
    "Primary Emergency Road backbone contains 2,562 target roots and yields 1,063.6 "
    "expected disconnected residents under Heavy rainfall. Expanding the target to "
    "2,977 Primary-plus-Secondary Emergency Road roots yields 992.7 residents, with "
    "community-frequency Spearman correlation 0.964 and top-30 burden overlap 90.0% "
    "relative to the primary definition. The former coast-inclusive boundary proxy "
    "is reported only as an audit comparator and reproduces the prior 1,121.7-resident "
    "result. Municipality-wide Yatsushiro assignments of 0.70 and 0.80 bound the revised "
    "primary result between 1,016.6 and 1,118.7 residents (−4.4% to +5.2% relative to the "
    "0.75 midpoint); community-frequency rank correlations remain 0.989–1.000 and Top-30 "
    "population-burden overlap is 93.3%. Under matched five-seed comparisons, the Low, "
    "Central, and High mappings (maximum section closure propensities 0.15, 0.30, and "
    "0.45) yield 351.4, 1,063.6, and 2,073.2 expected disconnected residents, respectively. "
    "Relative to Central, community-frequency rank correlation is 0.939 under Low and "
    "0.971 under High, while Top-30 population-burden overlap is 70.0% and 80.0%; 15 "
    "Top-30 communities are common to all three mappings (Appendix Table B12). The "
    "Low-to-High span is about 150 times the Central across-seed standard deviation, so "
    "mapping uncertainty dominates Monte Carlo variation in magnitude even though much "
    "of the priority ordering is retained."
)
EXPECTED_TITLE = (
    "Table B12. Heavy-scenario closure-mapping sensitivity under matched simulation seeds"
)
EXPECTED_HEADERS = [
    "Closure Mapping",
    "Maximum Section Closure Propensity",
    "Expected Disconnected Population",
    "Seed Minimum",
    "Seed Maximum",
    "Seed Standard Deviation",
    "Relative Change from Central",
    "Frequency Spearman Correlation with Central",
    "Top-30 Burden Overlap with Central",
    "Communities with Absolute Frequency Change >= 0.05",
    "Planning use",
]
WIDTHS = [0.52, 0.54, 0.65, 0.50, 0.50, 0.55, 0.58, 0.64, 0.62, 0.69, 0.81]


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def package_members(path: Path) -> dict[str, bytes]:
    with zipfile.ZipFile(path) as archive:
        bad = archive.testzip()
        if bad:
            raise RuntimeError(f"Corrupt DOCX member: {bad}")
        return {name: archive.read(name) for name in archive.namelist()}


def table_hash(table) -> str:
    return hashlib.sha256(table._tbl.xml.encode("utf-8")).hexdigest()


def table_text(table) -> list[list[str]]:
    return [[cell.text for cell in row.cells] for row in table.rows]


def bookmark_fingerprint(document: Document) -> str:
    parts: list[str] = []
    for element in document.element.body.iter():
        if element.tag in {qn("w:bookmarkStart"), qn("w:bookmarkEnd")}:
            parts.append(f"{element.tag}:{sorted(element.attrib.items())}")
    return hashlib.sha256("\n".join(parts).encode("utf-8")).hexdigest()


def workbook_values(path: Path) -> tuple[str, list[str], list[list[str]]]:
    workbook = load_workbook(path, data_only=True, read_only=True)
    sheet = workbook[workbook.sheetnames[0]]
    rows = list(sheet.iter_rows(values_only=True))
    title = str(rows[0][0])
    headers = [str(value) for value in rows[1][:11]]
    if title != EXPECTED_TITLE:
        raise RuntimeError(f"Unexpected Table B12 title: {title}")
    if headers != EXPECTED_HEADERS:
        raise RuntimeError(f"Unexpected Table B12 headers: {headers}")
    output: list[list[str]] = []
    for row in rows[2:5]:
        output.append(
            [
                str(row[0]),
                f"{float(row[1]):.2f}",
                f"{float(row[2]):,.1f}",
                f"{float(row[3]):,.1f}",
                f"{float(row[4]):,.1f}",
                f"{float(row[5]):.1f}",
                f"{100 * float(row[6]):.1f}%",
                f"{float(row[7]):.3f}",
                f"{100 * float(row[8]):.1f}%",
                f"{int(row[9])}",
                str(row[10]),
            ]
        )
    if [row[0] for row in output] != ["Low", "Central", "High"]:
        raise RuntimeError("Unexpected Table B12 mapping rows")
    return title, headers, output


def build_b12(document: Document, workbook_path: Path):
    title, headers, rows = workbook_values(workbook_path)
    table = build_table(document, title, headers, rows, WIDTHS, {0, 1})
    table.autofit = False
    for column, width in zip(table.columns, WIDTHS, strict=True):
        column.width = Inches(width)
        for cell in column.cells:
            cell.width = Inches(width)
    for cell in table.rows[1].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(5.0)
    for row in table.rows[2:]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(5.4)
    return table, [[title] * len(headers), headers, *rows]


def update(docx: Path, workbook_path: Path, dry_run: bool) -> dict[str, object]:
    before_hash = sha256(docx)
    before_members = package_members(docx)
    document = Document(docx)
    original_table_hashes = [table_hash(table) for table in document.tables]
    original_bookmarks = bookmark_fingerprint(document)

    matches = [p for p in document.paragraphs if visible_text(p) == OLD_PARAGRAPH]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one exact Appendix paragraph, found {len(matches)}")
    if any(table.cell(0, 0).text == EXPECTED_TITLE for table in document.tables):
        raise RuntimeError("Appendix Table B12 already exists")
    paragraph = matches[0]
    replace_visible_content(paragraph, NEW_PARAGRAPH)
    table, expected_table_text = build_b12(document, workbook_path)
    paragraph._p.addnext(table._tbl)

    with tempfile.NamedTemporaryFile(
        prefix=f".{docx.name}.", suffix=".tmp.docx", dir=docx.parent, delete=False
    ) as handle:
        staged = Path(handle.name)
    try:
        document.save(staged)
        verified = Document(staged)
        if sum(visible_text(p) == NEW_PARAGRAPH for p in verified.paragraphs) != 1:
            raise RuntimeError("Saved Appendix paragraph verification failed")
        table_matches = [
            table for table in verified.tables if table.cell(0, 0).text == EXPECTED_TITLE
        ]
        if len(table_matches) != 1:
            raise RuntimeError(f"Expected one saved Appendix Table B12, found {len(table_matches)}")
        checked = table_matches[0]
        if (len(checked.rows), len(checked.columns)) != (5, 11):
            raise RuntimeError("Appendix Table B12 shape verification failed")
        if table_text(checked) != expected_table_text:
            raise RuntimeError("Appendix Table B12 does not match the validated workbook")

        new_hash = table_hash(checked)
        unchanged_hashes = [
            value for value in [table_hash(table) for table in verified.tables] if value != new_hash
        ]
        if unchanged_hashes != original_table_hashes:
            raise RuntimeError("A pre-existing Appendix table changed")
        if bookmark_fingerprint(verified) != original_bookmarks:
            raise RuntimeError("Appendix bookmark fingerprint changed")

        after_members = package_members(staged)
        changed_members = sorted(
            name
            for name in set(before_members) | set(after_members)
            if before_members.get(name) != after_members.get(name)
        )
        unexpected = [
            name for name in changed_members if name not in {"word/document.xml", "docProps/core.xml"}
        ]
        if unexpected:
            raise RuntimeError(f"Protected Appendix package members changed: {unexpected}")

        receipt: dict[str, object] = {
            "status": "dry-run" if dry_run else "applied",
            "sha256_before": before_hash,
            "sha256_after": sha256(staged),
            "changed_package_members": changed_members,
            "paragraphs_replaced": 1,
            "table_b12_shape": [5, 11],
            "preexisting_table_count": len(original_table_hashes),
            "preexisting_tables_unchanged": True,
            "bookmarks_preserved": True,
            "backup": None,
        }
        if dry_run:
            return receipt

        backup_dir = docx.parent / ".kila-backups"
        backup_dir.mkdir(parents=True, exist_ok=True)
        stamp = dt.datetime.now().strftime("%Y%m%dT%H%M%S%f")
        backup = backup_dir / f"Appendix.before-r4c4.{stamp}.docx"
        shutil.copy2(docx, backup)
        if sha256(backup) != before_hash:
            raise RuntimeError("Appendix backup hash mismatch")
        os.replace(staged, docx)
        receipt["backup"] = str(backup)
        if sha256(docx) != receipt["sha256_after"]:
            raise RuntimeError("Final Appendix hash differs from validated staged file")
        return receipt
    finally:
        if staged.exists():
            staged.unlink()


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--docx", type=Path, default=Path("Rev/revision/Appendix.docx"))
    parser.add_argument(
        "--workbook",
        type=Path,
        default=Path(
            "data/exp/revision/reviewer-4-comment-4/"
            "Table_closure_mapping_policy_sensitivity.xlsx"
        ),
    )
    parser.add_argument("--apply", action="store_true")
    args = parser.parse_args()
    print(json.dumps(update(args.docx, args.workbook, dry_run=not args.apply), indent=2))


if __name__ == "__main__":
    main()
