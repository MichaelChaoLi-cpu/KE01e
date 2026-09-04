#!/usr/bin/env python3
"""Relocate Appendix Table B13 from portrait narrative pages to the landscape table section."""

from __future__ import annotations

import argparse
import datetime as dt
import hashlib
import json
import os
from pathlib import Path
import shutil
import tempfile
import zipfile

from docx import Document
from docx.oxml.ns import qn

from revision_update_appendix_r2c5 import page_break_paragraph
from revision_update_appendix_r2c6 import EXPECTED_TITLE


C1_TITLE = "Appendix Table C1. Municipality isolation and service-loss summary"
NOTE_PREFIX = "Notes: Heavy rainfall, Primary Emergency Road backbone"


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def members(path: Path) -> dict[str, bytes]:
    with zipfile.ZipFile(path) as archive:
        bad = archive.testzip()
        if bad:
            raise RuntimeError(f"Corrupt DOCX member: {bad}")
        return {name: archive.read(name) for name in archive.namelist()}


def bookmark_fingerprint(document: Document) -> str:
    values = []
    for element in document.element.body.iter():
        if element.tag in {qn("w:bookmarkStart"), qn("w:bookmarkEnd")}:
            values.append(f"{element.tag}:{sorted(element.attrib.items())}")
    return hashlib.sha256("\n".join(values).encode()).hexdigest()


def body_position(document: Document, element) -> int:
    return list(document.element.body).index(element)


def update(docx: Path, dry_run: bool) -> dict[str, object]:
    before_hash = sha256(docx)
    before_members = members(docx)
    document = Document(docx)
    before_bookmarks = bookmark_fingerprint(document)
    before_table_xml = [table._tbl.xml for table in document.tables]

    tables = [table for table in document.tables if table.cell(0, 0).text == EXPECTED_TITLE]
    notes = [paragraph for paragraph in document.paragraphs if paragraph.text.startswith(NOTE_PREFIX)]
    c1_titles = [paragraph for paragraph in document.paragraphs if paragraph.text == C1_TITLE]
    if len(tables) != 1 or len(notes) != 1 or len(c1_titles) != 1:
        raise RuntimeError(
            f"Expected one B13 table, note, and C1 title; found {len(tables)}, {len(notes)}, {len(c1_titles)}"
        )
    table = tables[0]
    note = notes[0]
    c1 = c1_titles[0]
    if body_position(document, table._tbl) > body_position(document, c1._p):
        raise RuntimeError("Table B13 is already in the landscape table section")

    document.element.body.remove(table._tbl)
    document.element.body.remove(note._p)
    c1._p.addprevious(page_break_paragraph())
    c1._p.addprevious(table._tbl)
    c1._p.addprevious(note._p)

    with tempfile.NamedTemporaryFile(
        prefix=f".{docx.name}.", suffix=".tmp.docx", dir=docx.parent, delete=False
    ) as handle:
        staged = Path(handle.name)
    try:
        document.save(staged)
        check = Document(staged)
        checked_tables = [
            candidate for candidate in check.tables if candidate.cell(0, 0).text == EXPECTED_TITLE
        ]
        checked_notes = [p for p in check.paragraphs if p.text.startswith(NOTE_PREFIX)]
        checked_c1 = [p for p in check.paragraphs if p.text == C1_TITLE]
        if len(checked_tables) != 1 or len(checked_notes) != 1 or len(checked_c1) != 1:
            raise RuntimeError("Saved Appendix relocation verification failed")
        if body_position(check, checked_tables[0]._tbl) > body_position(check, checked_c1[0]._p):
            raise RuntimeError("Saved Table B13 does not precede Table C1")
        before_b13_xml = table._tbl.xml
        before_other_xml = [xml for xml in before_table_xml if xml != before_b13_xml]
        after_b13_xml = checked_tables[0]._tbl.xml
        after_other_xml = [t._tbl.xml for t in check.tables if t.cell(0, 0).text != EXPECTED_TITLE]
        if after_b13_xml != before_b13_xml or after_other_xml != before_other_xml:
            raise RuntimeError("Table content or pre-existing table order changed during layout repair")
        if bookmark_fingerprint(check) != before_bookmarks:
            raise RuntimeError("Appendix bookmark fingerprint changed")

        after_members = members(staged)
        changed = sorted(
            name
            for name in set(before_members) | set(after_members)
            if before_members.get(name) != after_members.get(name)
        )
        unexpected = [name for name in changed if name not in {"word/document.xml", "docProps/core.xml"}]
        if unexpected:
            raise RuntimeError(f"Protected Appendix package members changed: {unexpected}")
        receipt: dict[str, object] = {
            "status": "dry-run" if dry_run else "applied",
            "sha256_before": before_hash,
            "sha256_after": sha256(staged),
            "changed_package_members": changed,
            "table_b13_shape": [len(checked_tables[0].rows), len(checked_tables[0].columns)],
            "table_content_unchanged": True,
            "bookmarks_preserved": True,
            "backup": None,
        }
        if dry_run:
            return receipt

        backup_dir = docx.parent / ".kila-backups"
        backup_dir.mkdir(parents=True, exist_ok=True)
        stamp = dt.datetime.now().strftime("%Y%m%dT%H%M%S%f")
        backup = backup_dir / f"Appendix.before-r2c6-table-reflow.{stamp}.docx"
        shutil.copy2(docx, backup)
        if sha256(backup) != before_hash:
            raise RuntimeError("Appendix backup hash mismatch")
        os.replace(staged, docx)
        receipt["backup"] = str(backup)
        return receipt
    finally:
        if staged.exists():
            staged.unlink()


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--docx", type=Path, default=Path("Rev/revision/Appendix.docx"))
    parser.add_argument("--apply", action="store_true")
    args = parser.parse_args()
    print(json.dumps(update(args.docx, dry_run=not args.apply), indent=2))


if __name__ == "__main__":
    main()
