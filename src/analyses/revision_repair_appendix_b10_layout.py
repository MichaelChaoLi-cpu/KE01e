#!/usr/bin/env python3
"""Repair only the approved Appendix Table B10 page-width layout."""

from __future__ import annotations

import argparse
import datetime as dt
import hashlib
import json
import os
from pathlib import Path
import shutil
import tempfile
import zipfile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches


TITLE = "Table B10. Emergency-road-backbone target sensitivity"
WIDTHS = [1.02, 0.54, 0.45, 0.54, 0.50, 0.60, 0.87, 0.90, 0.64, 0.64]


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def members(path: Path) -> dict[str, bytes]:
    with zipfile.ZipFile(path) as archive:
        bad = archive.testzip()
        if bad:
            raise RuntimeError(f"Corrupt DOCX member: {bad}")
        return {name: archive.read(name) for name in archive.namelist()}


def table_hash(table) -> str:
    return hashlib.sha256(table._tbl.xml.encode("utf-8")).hexdigest()


def find_b10(document: Document):
    matches = [
        table for table in document.tables
        if table.cell(0, 0).text == TITLE and len(table.columns) == 10
    ]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one B10 table, found {len(matches)}")
    return matches[0]


def run(docx: Path, dry_run: bool) -> dict[str, object]:
    before_hash = sha256(docx)
    before_members = members(docx)
    document = Document(docx)
    target = find_b10(document)
    other_hashes = [
        table_hash(table) for table in document.tables
        if table.cell(0, 0).text != TITLE
    ]

    target.autofit = False
    grid_columns = target._tbl.tblGrid.gridCol_lst
    if len(grid_columns) != len(WIDTHS):
        raise RuntimeError("B10 table-grid shape mismatch")
    for index, width in enumerate(WIDTHS):
        grid_columns[index].w = Inches(width)
        target.columns[index].width = Inches(width)
        for cell in target.columns[index].cells:
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            marker = tc_pr.find(qn("w:noWrap"))
            if marker is not None:
                tc_pr.remove(marker)

    with tempfile.NamedTemporaryFile(prefix=f".{docx.name}.", suffix=".tmp.docx", dir=docx.parent, delete=False) as handle:
        staged = Path(handle.name)
    try:
        document.save(staged)
        verified = Document(staged)
        checked = find_b10(verified)
        if (len(checked.rows), len(checked.columns)) != (11, 10):
            raise RuntimeError("B10 shape changed")
        checked_other_hashes = [table_hash(table) for table in verified.tables if table.cell(0, 0).text != TITLE]
        if checked_other_hashes != other_hashes:
            raise RuntimeError("A non-B10 table changed")
        after_members = members(staged)
        changed = sorted(
            name for name in set(before_members) | set(after_members)
            if before_members.get(name) != after_members.get(name)
        )
        forbidden = [name for name in changed if name.startswith("word/") and name != "word/document.xml"]
        if forbidden:
            raise RuntimeError(f"Protected Word members changed: {forbidden}")
        receipt: dict[str, object] = {
            "status": "dry-run" if dry_run else "applied",
            "sha256_before": before_hash,
            "sha256_after": sha256(staged),
            "changed_package_members": changed,
            "b10_total_width_inches": sum(WIDTHS),
            "non_b10_tables_unchanged": True,
            "backup": None,
        }
        if dry_run:
            return receipt
        backup_dir = docx.parent / ".kila-backups"
        backup_dir.mkdir(parents=True, exist_ok=True)
        stamp = dt.datetime.now().strftime("%Y%m%dT%H%M%S%f")
        backup = backup_dir / f"Appendix.before-b10-layout-repair.{stamp}.docx"
        shutil.copy2(docx, backup)
        if sha256(backup) != before_hash:
            raise RuntimeError("Backup hash mismatch")
        os.replace(staged, docx)
        receipt["backup"] = str(backup)
        return receipt
    finally:
        if staged.exists():
            staged.unlink()


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--docx", type=Path, default=Path("Rev/revision/Appendix.docx"))
    parser.add_argument("--apply", action="store_true")
    args = parser.parse_args()
    print(json.dumps(run(args.docx, dry_run=not args.apply), indent=2))


if __name__ == "__main__":
    main()
