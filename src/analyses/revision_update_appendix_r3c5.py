#!/usr/bin/env python3
"""Apply the approved Reviewer 3 Comment 5 Appendix B8 update transactionally."""

from __future__ import annotations

import argparse
import datetime as dt
import hashlib
import json
import os
from pathlib import Path
import shutil
import tempfile
import zipfile

import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from revision_update_appendix_r2c5 import (
    build_table,
    insert_paragraph_after,
    page_break_paragraph,
    paragraph_with_exact_text,
)


ANCHOR = (
    "Simulation-size and network-definition checks bound the isolation results. "
    "Changing the number of draws from 500 to 2,000 with one common seed produces "
    "a 0.006 difference in the 95th-percentile community isolation frequency. "
    "Alternative external-road targets place Heavy-scenario expected isolation "
    "between 1,044 and 1,108 residents, compared with the five-seed Central mean "
    "of 1,106.9 under the primary target. Municipality-wide Yatsushiro assignments "
    "of 0.70 and 0.80 bound the result between 1,041 and 1,193, while alternative "
    "closure mappings produce the wider range of 343–2,309 residents. The closure "
    "mapping is therefore a larger source of uncertainty than Monte Carlo "
    "convergence or the tested gateway definitions."
)
NARRATIVE = (
    "Appendix Table B8 isolates spatial closure dependence while retaining each "
    "candidate section's central marginal closure propensity. The independent "
    "implementation exactly reproduces the existing five-seed simulation. Under "
    "the broad strong setting, mean expected isolated population changes by +9.2%, "
    "+15.7%, and −10.5% for Moderate, Heavy, and Extreme rainfall, respectively, "
    "and the corresponding per-draw 95th percentiles change by +8.7%, +15.8%, and "
    "−5.3%. Community-frequency rank correlations remain high, but lower top-30 "
    "overlap under Heavy rainfall shows that local preparedness priorities are more "
    "dependence-sensitive than the broad geographic ordering. Because the scale and "
    "correlation settings are not calibrated, the table reports sensitivity bounds "
    "rather than alternative forecasts."
)
TITLE = "Table B8. Spatial closure-dependence sensitivity"


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def package_members(path: Path) -> dict[str, bytes]:
    with zipfile.ZipFile(path) as archive:
        bad = archive.testzip()
        if bad:
            raise RuntimeError(f"Corrupt DOCX member: {bad}")
        return {name: archive.read(name) for name in archive.namelist()}


def table_xml_hashes(document: Document) -> list[str]:
    return [hashlib.sha256(table._tbl.xml.encode("utf-8")).hexdigest() for table in document.tables]


def paragraph_visible_text(paragraph) -> str:
    visible_tags = {qn("w:t"), qn("w:delText"), qn("m:t")}
    return "".join(
        element.text or ""
        for element in paragraph._p.iter()
        if element.tag in visible_tags
    )


def visible_paragraph_with_exact_text(document: Document, text: str):
    matches = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph_visible_text(paragraph) == text
    ]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one exact visible paragraph match, found {len(matches)}")
    return matches[0]


def fmt_change(value: float) -> str:
    if abs(value) < 0.00005:
        return "0.0%"
    return f"{value:+.1%}".replace("-", "−")


def table_rows(summary_path: Path) -> list[list[str]]:
    frame = pd.read_csv(summary_path)
    if len(frame) != 15:
        raise RuntimeError(f"Expected 15 sensitivity rows, found {len(frame)}")
    independent = (
        frame.loc[frame["dependence_setting"] == "Independent"]
        .set_index("scenario")
    )
    rows: list[list[str]] = []
    for row in frame.itertuples(index=False):
        reference = independent.loc[row.scenario]
        mean_change = row.expected_isolated_population_mean / reference.expected_isolated_population_mean - 1
        p95_change = row.draw_isolated_population_p95 / reference.draw_isolated_population_p95 - 1
        cluster = (
            "None; 0.00"
            if row.dependence_setting == "Independent"
            else f"{row.cluster_scale_km:.0f} km; {row.rho:.2f}"
        )
        rows.append(
            [
                row.scenario,
                row.dependence_setting,
                cluster,
                f"{row.expected_isolated_population_mean:,.1f} / "
                f"{row.expected_isolated_older_population_mean:,.1f}",
                fmt_change(mean_change),
                f"{row.draw_isolated_population_p95:,.1f}",
                fmt_change(p95_change),
                f"{row.community_frequency_spearman_vs_independent:.3f}",
                f"{100 * row.top30_burden_overlap_vs_independent:.1f}%",
                f"{int(row.communities_abs_frequency_change_ge_0_05):,}",
            ]
        )
    return rows


def apply_update(docx: Path, summary_path: Path, dry_run: bool) -> dict[str, object]:
    before_hash = sha256(docx)
    before_members = package_members(docx)
    document = Document(docx)
    original_table_hashes = table_xml_hashes(document)
    if any(paragraph.text == NARRATIVE for paragraph in document.paragraphs):
        raise RuntimeError("Appendix B8 narrative already exists")
    if any(table.cell(0, 0).text == TITLE for table in document.tables):
        raise RuntimeError("Appendix Table B8 already exists")

    anchor = visible_paragraph_with_exact_text(document, ANCHOR)
    insert_paragraph_after(anchor._p, NARRATIVE, "Normal")

    headers = [
        "Rainfall\nscenario",
        "Dependence\nsetting",
        "Cluster scale;\nrho",
        "Expected isolated\npopulation\n(total / age 65+)",
        "Mean change\nvs independent",
        "Per-draw\nP95",
        "P95 change\nvs independent",
        "Community-frequency\nSpearman",
        "Top-30 burden\noverlap",
        "Communities with\nabsolute frequency\nchange >= 0.05",
    ]
    table = build_table(
        document,
        TITLE,
        headers,
        table_rows(summary_path),
        [0.78, 0.93, 0.80, 1.15, 0.85, 0.72, 0.82, 0.92, 0.78, 1.05],
        {2, 3, 4, 5, 6, 7, 8, 9},
    )
    c1_title = paragraph_with_exact_text(
        document, "Appendix Table C1. Municipality isolation and service-loss summary"
    )._p
    c1_title.addprevious(page_break_paragraph())
    c1_title.addprevious(table._tbl)
    c1_title.addprevious(OxmlElement("w:p"))

    with tempfile.NamedTemporaryFile(
        prefix=f".{docx.name}.", suffix=".tmp.docx", dir=docx.parent, delete=False
    ) as handle:
        temp_path = Path(handle.name)
    try:
        document.save(temp_path)
        verified = Document(temp_path)
        texts = [paragraph.text for paragraph in verified.paragraphs]
        if texts.count(NARRATIVE) != 1:
            raise RuntimeError("Appendix B8 narrative verification failed")
        matches = [table for table in verified.tables if table.cell(0, 0).text == TITLE]
        if len(matches) != 1:
            raise RuntimeError(f"Expected one Appendix B8 table, found {len(matches)}")
        check_table = matches[0]
        if (len(check_table.rows), len(check_table.columns)) != (17, 10):
            raise RuntimeError("Appendix B8 table shape verification failed")

        new_table_hash = hashlib.sha256(check_table._tbl.xml.encode("utf-8")).hexdigest()
        unchanged_hashes = [value for value in table_xml_hashes(verified) if value != new_table_hash]
        if unchanged_hashes != original_table_hashes:
            raise RuntimeError("A pre-existing Appendix table changed")

        after_members = package_members(temp_path)
        changed_members = sorted(
            name
            for name in set(before_members) | set(after_members)
            if before_members.get(name) != after_members.get(name)
        )
        forbidden = [
            name
            for name in changed_members
            if name.startswith("word/") and name != "word/document.xml"
        ]
        if forbidden:
            raise RuntimeError(f"Protected Word package members changed: {forbidden}")

        receipt: dict[str, object] = {
            "status": "dry-run" if dry_run else "applied",
            "sha256_before": before_hash,
            "sha256_after": sha256(temp_path),
            "changed_package_members": changed_members,
            "table_b8_shape": [17, 10],
            "preexisting_table_count": len(original_table_hashes),
            "preexisting_tables_unchanged": True,
            "backup": None,
        }
        if dry_run:
            return receipt

        backup_dir = docx.parent / ".kila-backups"
        backup_dir.mkdir(parents=True, exist_ok=True)
        stamp = dt.datetime.now().strftime("%Y%m%dT%H%M%S%f")
        backup = backup_dir / f"{docx.stem}.{stamp}.reviewer-3-comment-5.part-06.docx"
        shutil.copy2(docx, backup)
        if sha256(backup) != before_hash:
            raise RuntimeError("Appendix backup hash mismatch")
        os.replace(temp_path, docx)
        receipt["backup"] = str(backup)
        if sha256(docx) != receipt["sha256_after"]:
            raise RuntimeError("Final Appendix hash differs from validated staged file")
        return receipt
    finally:
        if temp_path.exists():
            temp_path.unlink()


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--docx", type=Path, required=True)
    parser.add_argument("--summary", type=Path, required=True)
    parser.add_argument("--dry-run", action="store_true")
    args = parser.parse_args()
    print(json.dumps(apply_update(args.docx, args.summary, args.dry_run), indent=2))


if __name__ == "__main__":
    main()
