#!/usr/bin/env python3
"""Insert the approved Reviewer 1 Comment 4 comparator clarification."""

from __future__ import annotations

import argparse
import copy
import datetime as dt
import hashlib
import json
import os
from pathlib import Path
import shutil
import tempfile
import zipfile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree


ANCHOR = (
    "Appendix Table B3 shows that protected population increases with budget under "
    "Conservative, Central, and Optimistic assumptions. At the maximum budget of "
    "269.13 relative planning units, benefits are 31.2, 62.3, and 99.6 residents, "
    "while selection overlap with the Central portfolio is 84.0% in the Conservative "
    "setting and 80.6% in the Optimistic setting. Appendix Table B4 shows that the "
    "Central assigned-action and equal-cost consequence rankings both protect 62.3 "
    "residents, whereas hazard-only, emergency-route-only, and road-class-only rules "
    "protect 0.5, 0.3, and 0.0 residents, respectively. The comparison supports "
    "consequence-aware screening but not superiority over the equivalent consequence "
    "benchmark."
)

INSERTED = (
    "The comparator labels denote four fixed heuristics. Hazard only ranks Road "
    "Disruption Score; emergency route only and road class only prioritize the named "
    "binary or class criterion and use Road Disruption Score to break ties; equal-cost "
    "consequence ranks the same consequence-effect-to-cost ratio under each setting. "
    "Under the declared action effects and global cost multipliers, the median ratio "
    "used by the assigned-action score equals the Central ratio for all three actions, "
    "so the two Central orders and all seven budget rows are identical. No formal "
    "multi-criteria decision model was tested because the study did not have "
    "stakeholder-elicited criterion weights or a validated cross-criterion utility "
    "function."
)

NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def read_package(path: Path) -> tuple[list[zipfile.ZipInfo], dict[str, bytes]]:
    with zipfile.ZipFile(path) as archive:
        bad = archive.testzip()
        if bad:
            raise RuntimeError(f"Corrupt DOCX member: {bad}")
        return archive.infolist(), {name: archive.read(name) for name in archive.namelist()}


def write_package(path: Path, infos: list[zipfile.ZipInfo], files: dict[str, bytes]) -> None:
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as archive:
        for info in infos:
            archive.writestr(info, files[info.filename])


def paragraph_text(paragraph: etree._Element) -> str:
    return "".join(paragraph.xpath(".//w:t/text()", namespaces=NS))


def bookmark_fingerprint(root: etree._Element) -> str:
    parts: list[str] = []
    for node in root.xpath(".//w:bookmarkStart | .//w:bookmarkEnd", namespaces=NS):
        parts.append(f"{node.tag}:{sorted(node.attrib.items())}")
    return hashlib.sha256("\n".join(parts).encode("utf-8")).hexdigest()


def table_hashes(document: Document) -> list[str]:
    return [hashlib.sha256(table._tbl.xml.encode("utf-8")).hexdigest() for table in document.tables]


def build_paragraph(anchor: etree._Element) -> etree._Element:
    paragraph = OxmlElement("w:p")
    ppr = anchor.find(qn("w:pPr"))
    if ppr is not None:
        paragraph.append(copy.deepcopy(ppr))
    run = OxmlElement("w:r")
    source_runs = anchor.xpath("./w:r[w:t]", namespaces=NS)
    if source_runs:
        rpr = source_runs[-1].find(qn("w:rPr"))
        if rpr is not None:
            run.append(copy.deepcopy(rpr))
    text = OxmlElement("w:t")
    text.text = INSERTED
    run.append(text)
    paragraph.append(run)
    return paragraph


def update(docx: Path, dry_run: bool) -> dict[str, object]:
    before_hash = sha256(docx)
    infos, files = read_package(docx)
    original_files = dict(files)
    original_root = etree.fromstring(files["word/document.xml"])
    original_bookmarks = bookmark_fingerprint(original_root)
    original_tables = table_hashes(Document(docx))

    paragraphs = original_root.xpath(".//w:body/w:p", namespaces=NS)
    matches = [paragraph for paragraph in paragraphs if paragraph_text(paragraph) == ANCHOR]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one exact Appendix anchor, found {len(matches)}")
    if sum(paragraph_text(paragraph) == INSERTED for paragraph in paragraphs):
        raise RuntimeError("Approved Appendix comparator paragraph already exists")

    anchor = matches[0]
    parent = anchor.getparent()
    parent.insert(parent.index(anchor) + 1, build_paragraph(anchor))
    files["word/document.xml"] = etree.tostring(
        original_root, xml_declaration=True, encoding="UTF-8", standalone=True
    )

    with tempfile.NamedTemporaryFile(
        prefix=f".{docx.name}.", suffix=".tmp.docx", dir=docx.parent, delete=False
    ) as handle:
        staged = Path(handle.name)
    try:
        write_package(staged, infos, files)
        _, staged_files = read_package(staged)
        staged_root = etree.fromstring(staged_files["word/document.xml"])
        saved_paragraphs = staged_root.xpath(".//w:body/w:p", namespaces=NS)
        if sum(paragraph_text(paragraph) == INSERTED for paragraph in saved_paragraphs) != 1:
            raise RuntimeError("Saved Appendix paragraph verification failed")
        if bookmark_fingerprint(staged_root) != original_bookmarks:
            raise RuntimeError("Appendix bookmark fingerprint changed")
        verified = Document(staged)
        if len(verified.tables) != len(original_tables):
            raise RuntimeError("Appendix table count changed")
        if table_hashes(verified) != original_tables:
            raise RuntimeError("A pre-existing Appendix table changed")

        changed_members = sorted(
            name for name in files if staged_files.get(name) != original_files.get(name)
        )
        if changed_members != ["word/document.xml"]:
            raise RuntimeError(f"Protected Appendix package members changed: {changed_members}")

        receipt: dict[str, object] = {
            "status": "dry-run" if dry_run else "applied",
            "sha256_before": before_hash,
            "sha256_after": sha256(staged),
            "changed_package_members": changed_members,
            "paragraphs_inserted": 1,
            "preexisting_table_count": len(original_tables),
            "preexisting_tables_unchanged": True,
            "bookmarks_preserved": True,
            "backup": None,
        }
        if dry_run:
            return receipt

        backup_dir = docx.parent / ".kila-backups"
        backup_dir.mkdir(parents=True, exist_ok=True)
        stamp = dt.datetime.now().strftime("%Y%m%dT%H%M%S%f")
        backup = backup_dir / f"Appendix.before-r1c4.{stamp}.docx"
        shutil.copy2(docx, backup)
        if sha256(backup) != before_hash:
            raise RuntimeError("Appendix backup hash mismatch")
        os.replace(staged, docx)
        receipt["backup"] = str(backup)
        if sha256(docx) != receipt["sha256_after"]:
            raise RuntimeError("Final Appendix hash differs from validated staged file")
        return receipt
    finally:
        if staged.exists():
            staged.unlink()


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--docx", type=Path, required=True)
    parser.add_argument("--dry-run", action="store_true")
    args = parser.parse_args()
    print(json.dumps(update(args.docx, args.dry_run), indent=2))


if __name__ == "__main__":
    main()
