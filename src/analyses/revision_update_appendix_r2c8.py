#!/usr/bin/env python3
"""Apply the approved Reviewer 2 Comment 8 Appendix B9/C1 update."""

from __future__ import annotations

import argparse
import datetime as dt
import hashlib
import json
import os
from pathlib import Path
import shutil
import tempfile
import zipfile

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches

from revision_update_appendix_r2c5 import (
    build_table,
    insert_paragraph_after,
    set_cell_text,
    set_row_cant_split,
)


TITLE = "Appendix Table B9. Service-destination estimand and rerouting comparison"
NARRATIVE = (
    "Appendix Table B9 compares the implemented any-same-class reachability "
    "estimand with a restrictive fixed-baseline-destination continuity comparator "
    "using identical Heavy-scenario closure draws. The any-same-class branch exactly "
    "reproduces all 20 production seed-by-service arrays. Fixed-destination loss is "
    "larger for every service class because it counts communities that lose the "
    "baseline-nearest facility even when another same-class destination remains "
    "reachable. The difference quantifies rerouting benefit, not observed facility "
    "substitution, and does not represent capacity or operability at the alternative "
    "destination."
)
SERVICE_SECTION_ANCHOR = (
    "Service sensitivity retains only baseline-reachable community-destination pairs "
    "and does not convert pre-existing disconnection into disruption-induced loss."
)
SERVICE_ORDER = ["Shelter", "Fire service", "Municipal facility", "Emergency water"]
B9_HEADERS = [
    "Service",
    "Sources\nresolved / total",
    "Road\nattached",
    "Eligible\ncommunities",
    "Eligible\npopulation",
    "Any-class\nloss mean",
    "Any-class\nseed range",
    "Fixed loss\nmean",
    "Fixed loss\nseed range",
    "Rerouting\nbenefit",
    "Benefit /\nfixed loss",
    "Communities\nbenefiting",
]
B9_WIDTHS = [0.70, 0.50, 0.40, 0.52, 0.56, 0.54, 0.52, 0.54, 0.52, 0.54, 0.46, 0.48]
OLD_HEADERS = {
    "Heavy Shelter Loss Population (Baseline-Reachable)": (
        "Heavy Shelter Loss Population (Any Same-Class Facility)"
    ),
    "Heavy Emergency-Water Sensitivity Loss Population (10/36 Geolocated)": (
        "Heavy Emergency-Water Sensitivity Loss Population "
        "(Any of 10/36 Geolocated Facilities)"
    ),
    "Heavy Fire service Loss Population (Baseline-Reachable)": (
        "Heavy Fire service Loss Population (Any Same-Class Facility)"
    ),
    "Heavy Municipal facility Loss Population (Baseline-Reachable)": (
        "Heavy Municipal facility Loss Population (Any Same-Class Facility)"
    ),
}


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def package_members(path: Path) -> dict[str, bytes]:
    with zipfile.ZipFile(path) as archive:
        bad = archive.testzip()
        if bad:
            raise RuntimeError(f"Corrupt DOCX member: {bad}")
        return {name: archive.read(name) for name in archive.namelist()}


def table_hash(table) -> str:
    return hashlib.sha256(table._tbl.xml.encode("utf-8")).hexdigest()


def table_text(table) -> list[list[str]]:
    return [[cell.text for cell in row.cells] for row in table.rows]


def table_rows(summary_path: Path) -> list[list[str]]:
    frame = pd.read_csv(summary_path).set_index("Service Class").loc[SERVICE_ORDER]
    rows: list[list[str]] = []
    for service, row in frame.iterrows():
        label = "Emergency water (conditional)" if service == "Emergency water" else service
        rows.append(
            [
                label,
                f"{int(row['Resolved Source Facilities']):,} / "
                f"{int(row['Source Facilities Total']):,}",
                f"{int(row['Road-Attached Facilities']):,}",
                f"{int(row['Baseline-Eligible Communities']):,}",
                f"{float(row['Baseline-Eligible Population']):,.1f}",
                f"{float(row['Any-Same-Class Loss Population Mean']):,.1f}",
                f"{float(row['Any-Same-Class Loss Population Min']):,.1f}-"
                f"{float(row['Any-Same-Class Loss Population Max']):,.1f}",
                f"{float(row['Fixed-Destination Loss Population Mean']):,.1f}",
                f"{float(row['Fixed-Destination Loss Population Min']):,.1f}-"
                f"{float(row['Fixed-Destination Loss Population Max']):,.1f}",
                f"{float(row['Rerouting Benefit Population Mean']):,.1f}",
                f"{100 * float(row['Rerouting Benefit Share of Fixed Loss']):.1f}%",
                f"{int(row['Communities with Positive Rerouting Benefit']):,}",
            ]
        )
    return rows


def find_c1(document: Document):
    matches = [
        table
        for table in document.tables
        if (len(table.rows), len(table.columns)) == (50, 16)
        and table.cell(0, 0).text == "Admin Area Code"
    ]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one Appendix Table C1, found {len(matches)}")
    return matches[0]


def replace_header_text(cell, old: str, new: str) -> None:
    if cell.text != old:
        raise RuntimeError(f"Appendix C1 header mismatch: expected {old!r}, found {cell.text!r}")
    paragraph = cell.paragraphs[0]
    if len(paragraph.runs) != 1:
        raise RuntimeError(f"Appendix C1 header {old!r} is not a single safe run")
    paragraph.runs[0].text = new


def format_b9(table) -> None:
    """Fit the approved 12-column audit table within a portrait text block."""
    table.autofit = False
    layout = table._tbl.tblPr.find(qn("w:tblLayout"))
    if layout is not None:
        layout.set(qn("w:type"), "fixed")
    for column, width in zip(table.columns, B9_WIDTHS):
        column.width = Inches(width)
        for cell in column.cells:
            cell.width = Inches(width)
            no_wrap = cell._tc.get_or_add_tcPr().find(qn("w:noWrap"))
            if no_wrap is not None:
                cell._tc.get_or_add_tcPr().remove(no_wrap)

    set_cell_text(
        table.cell(0, 0),
        TITLE,
        bold=True,
        size=9.5,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        color="1F4E78",
        shading="D9EAF7",
    )
    for column, header in enumerate(B9_HEADERS):
        set_cell_text(
            table.cell(1, column),
            header,
            bold=True,
            size=5.1,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            color="FFFFFF",
            shading="1F4E78",
        )
    for row_index, values in enumerate(table_text(table)[2:], start=2):
        for column, value in enumerate(values):
            set_cell_text(
                table.cell(row_index, column),
                value,
                bold=False,
                size=5.0,
                align=(
                    WD_ALIGN_PARAGRAPH.LEFT
                    if column == 0
                    else WD_ALIGN_PARAGRAPH.CENTER
                ),
                shading="F4F7F9" if row_index % 2 == 0 else "FFFFFF",
            )
        set_row_cant_split(table.rows[row_index])


def apply_update(docx: Path, summary_path: Path, dry_run: bool) -> dict[str, object]:
    before_hash = sha256(docx)
    before_members = package_members(docx)
    document = Document(docx)
    if any(paragraph.text == NARRATIVE for paragraph in document.paragraphs):
        raise RuntimeError("Appendix B9 narrative already exists")
    if any(table.cell(0, 0).text == TITLE for table in document.tables):
        raise RuntimeError("Appendix Table B9 already exists")

    c1 = find_c1(document)
    c1_index = next(
        index for index, table in enumerate(document.tables) if table._tbl is c1._tbl
    )
    c1_before = table_text(c1)
    protected_before = [
        table_hash(table)
        for index, table in enumerate(document.tables)
        if index != c1_index
    ]

    anchors = [
        paragraph
        for paragraph in document.paragraphs
        if SERVICE_SECTION_ANCHOR in paragraph.text
    ]
    if len(anchors) != 1:
        raise RuntimeError(f"Expected one service-sensitivity anchor, found {len(anchors)}")
    narrative_element = insert_paragraph_after(anchors[0]._p, NARRATIVE, "Normal")

    b9 = build_table(
        document,
        TITLE,
        B9_HEADERS,
        table_rows(summary_path),
        B9_WIDTHS,
        set(),
    )
    format_b9(b9)
    narrative_element.addnext(b9._tbl)

    current_headers = [cell.text for cell in c1.rows[0].cells]
    positions = {header: current_headers.index(header) for header in OLD_HEADERS}
    if len(set(positions.values())) != 4:
        raise RuntimeError("Appendix C1 service headers are not uniquely identifiable")
    for old, new in OLD_HEADERS.items():
        replace_header_text(c1.cell(0, positions[old]), old, new)

    with tempfile.NamedTemporaryFile(
        prefix=f".{docx.name}.", suffix=".tmp.docx", dir=docx.parent, delete=False
    ) as handle:
        staged = Path(handle.name)
    try:
        document.save(staged)
        verified = Document(staged)
        narrative_matches = [p for p in verified.paragraphs if p.text == NARRATIVE]
        if len(narrative_matches) != 1:
            raise RuntimeError("Appendix B9 narrative verification failed")
        b9_matches = [t for t in verified.tables if t.cell(0, 0).text == TITLE]
        if len(b9_matches) != 1:
            raise RuntimeError(f"Expected one Appendix Table B9, found {len(b9_matches)}")
        checked_b9 = b9_matches[0]
        if (len(checked_b9.rows), len(checked_b9.columns)) != (6, 12):
            raise RuntimeError("Appendix Table B9 shape verification failed")
        if table_text(checked_b9)[2:] != table_rows(summary_path):
            raise RuntimeError("Appendix Table B9 values do not match the validated summary")

        checked_c1 = find_c1(verified)
        c1_after = table_text(checked_c1)
        for row_index in range(1, len(c1_before)):
            if c1_after[row_index] != c1_before[row_index]:
                raise RuntimeError("An Appendix C1 data value changed")
        expected_headers = list(c1_before[0])
        for old, new in OLD_HEADERS.items():
            expected_headers[expected_headers.index(old)] = new
        if c1_after[0] != expected_headers:
            raise RuntimeError("Appendix C1 header-only update verification failed")

        protected_after = [
            table_hash(table)
            for table in verified.tables
            if table._tbl is not checked_c1._tbl and table.cell(0, 0).text != TITLE
        ]
        if protected_after != protected_before:
            raise RuntimeError("A pre-existing Appendix table outside C1 changed")

        after_members = package_members(staged)
        changed_members = sorted(
            name
            for name in set(before_members) | set(after_members)
            if before_members.get(name) != after_members.get(name)
        )
        allowed = {"word/document.xml", "docProps/core.xml"}
        unexpected = [name for name in changed_members if name not in allowed]
        if unexpected:
            raise RuntimeError(f"Protected DOCX package members changed: {unexpected}")

        receipt: dict[str, object] = {
            "status": "dry-run" if dry_run else "applied",
            "sha256_before": before_hash,
            "sha256_after": sha256(staged),
            "changed_package_members": changed_members,
            "table_b9_shape": [6, 12],
            "table_b9_rows_match_validated_summary": True,
            "table_c1_data_unchanged": True,
            "table_c1_headers_renamed": list(OLD_HEADERS.values()),
            "other_tables_unchanged": True,
            "backup": None,
        }
        if dry_run:
            return receipt

        backup_dir = docx.parent / ".kila-backups"
        backup_dir.mkdir(parents=True, exist_ok=True)
        stamp = dt.datetime.now().strftime("%Y%m%dT%H%M%S%f")
        backup = backup_dir / f"{docx.stem}.{stamp}.reviewer-2-comment-8.part-08.docx"
        shutil.copy2(docx, backup)
        if sha256(backup) != before_hash:
            raise RuntimeError("Appendix backup hash mismatch")
        os.replace(staged, docx)
        receipt["backup"] = str(backup)
        if sha256(docx) != receipt["sha256_after"]:
            raise RuntimeError("Final Appendix hash differs from the validated staged file")
        return receipt
    finally:
        if staged.exists():
            staged.unlink()


def apply_layout_repair(docx: Path, summary_path: Path, dry_run: bool) -> dict[str, object]:
    """Repair only the approved B9 layout after rendered-page inspection."""
    before_hash = sha256(docx)
    before_members = package_members(docx)
    document = Document(docx)
    matches = [table for table in document.tables if table.cell(0, 0).text == TITLE]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one Appendix Table B9, found {len(matches)}")
    b9 = matches[0]
    if table_text(b9)[2:] != table_rows(summary_path):
        raise RuntimeError("Existing Appendix Table B9 values do not match validated summary")
    protected_before = [
        table_hash(table) for table in document.tables if table._tbl is not b9._tbl
    ]
    c1_before = table_text(find_c1(document))
    format_b9(b9)

    with tempfile.NamedTemporaryFile(
        prefix=f".{docx.name}.", suffix=".tmp.docx", dir=docx.parent, delete=False
    ) as handle:
        staged = Path(handle.name)
    try:
        document.save(staged)
        verified = Document(staged)
        checked = [table for table in verified.tables if table.cell(0, 0).text == TITLE]
        if len(checked) != 1 or table_text(checked[0])[2:] != table_rows(summary_path):
            raise RuntimeError("Reformatted Appendix Table B9 failed value verification")
        if [cell.text for cell in checked[0].rows[1].cells] != B9_HEADERS:
            raise RuntimeError("Reformatted Appendix Table B9 headers are incorrect")
        if table_text(find_c1(verified)) != c1_before:
            raise RuntimeError("Appendix C1 changed during B9 layout repair")
        protected_after = [
            table_hash(table)
            for table in verified.tables
            if table.cell(0, 0).text != TITLE
        ]
        if protected_after != protected_before:
            raise RuntimeError("A protected Appendix table changed during B9 layout repair")
        after_members = package_members(staged)
        changed_members = sorted(
            name
            for name in set(before_members) | set(after_members)
            if before_members.get(name) != after_members.get(name)
        )
        unexpected = [
            name
            for name in changed_members
            if name not in {"word/document.xml", "docProps/core.xml"}
        ]
        if unexpected:
            raise RuntimeError(f"Protected DOCX package members changed: {unexpected}")
        receipt: dict[str, object] = {
            "status": "dry-run" if dry_run else "applied",
            "sha256_before": before_hash,
            "sha256_after": sha256(staged),
            "changed_package_members": changed_members,
            "table_b9_values_unchanged": True,
            "table_c1_unchanged": True,
            "other_tables_unchanged": True,
            "backup": None,
        }
        if dry_run:
            return receipt
        backup_dir = docx.parent / ".kila-backups"
        backup_dir.mkdir(parents=True, exist_ok=True)
        stamp = dt.datetime.now().strftime("%Y%m%dT%H%M%S%f")
        backup = backup_dir / f"{docx.stem}.{stamp}.reviewer-2-comment-8.b9-layout.docx"
        shutil.copy2(docx, backup)
        if sha256(backup) != before_hash:
            raise RuntimeError("Appendix layout-repair backup hash mismatch")
        os.replace(staged, docx)
        receipt["backup"] = str(backup)
        return receipt
    finally:
        if staged.exists():
            staged.unlink()


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--docx", type=Path, required=True)
    parser.add_argument("--summary", type=Path, required=True)
    parser.add_argument("--dry-run", action="store_true")
    parser.add_argument("--repair-existing", action="store_true")
    args = parser.parse_args()
    function = apply_layout_repair if args.repair_existing else apply_update
    print(json.dumps(function(args.docx, args.summary, args.dry_run), indent=2))


if __name__ == "__main__":
    main()
