#!/usr/bin/env python3
"""Apply the approved R3C6 supplemental Appendix paragraph update."""

from __future__ import annotations

import argparse
from copy import deepcopy
import datetime as dt
import hashlib
import json
import os
from pathlib import Path
import shutil
import tempfile
import zipfile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OLD = (
    "Appendix Table B3 shows that protected population increases with budget under "
    "Conservative, Central, and Optimistic assumptions. At the maximum budget of "
    "270.28 relative planning units, benefits are 31.0, 63.2, and 101.1 residents, "
    "while selection overlap with the Central portfolio is 84.0% in the Conservative "
    "setting and 80.6% in the Optimistic setting. Appendix Table B4 shows that the "
    "Central assigned-action and equal-cost consequence rankings both protect 63.2 "
    "residents, whereas hazard-only, emergency-route-only, and road-class-only rules "
    "protect 0.5, 0.4, and 0.7, respectively. The comparison supports consequence-aware "
    "screening but not superiority over the equivalent consequence benchmark."
)
NEW = (
    "Appendix Table B3 shows that protected population increases with budget under "
    "Conservative, Central, and Optimistic assumptions. At the maximum budget of "
    "269.13 relative planning units, benefits are 31.2, 62.3, and 99.6 residents, "
    "while selection overlap with the Central portfolio is 84.0% in the Conservative "
    "setting and 80.6% in the Optimistic setting. Appendix Table B4 shows that the "
    "Central assigned-action and equal-cost consequence rankings both protect 62.3 "
    "residents, whereas hazard-only, emergency-route-only, and road-class-only rules "
    "protect 0.5, 0.3, and 0.0 residents, respectively. The comparison supports "
    "consequence-aware screening but not superiority over the equivalent consequence "
    "benchmark."
)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def package_members(path: Path) -> dict[str, bytes]:
    with zipfile.ZipFile(path) as archive:
        bad = archive.testzip()
        if bad:
            raise RuntimeError(f"Corrupt DOCX member: {bad}")
        return {name: archive.read(name) for name in archive.namelist()}


def table_hashes(document: Document) -> list[str]:
    return [hashlib.sha256(table._tbl.xml.encode("utf-8")).hexdigest() for table in document.tables]


def visible_text(paragraph) -> str:
    tags = {qn("w:t"), qn("w:delText"), qn("m:t")}
    return "".join(element.text or "" for element in paragraph._p.iter() if element.tag in tags)


def replace_text(paragraph, text_value: str) -> None:
    p = paragraph._p
    p_pr = p.find(qn("w:pPr"))
    first_run = p.find(qn("w:r"))
    r_pr = None
    if first_run is not None and first_run.find(qn("w:rPr")) is not None:
        r_pr = deepcopy(first_run.find(qn("w:rPr")))
    for child in list(p):
        if child is not p_pr:
            p.remove(child)
    run = OxmlElement("w:r")
    if r_pr is not None:
        run.append(r_pr)
    text = OxmlElement("w:t")
    text.text = text_value
    run.append(text)
    p.append(run)


def update(docx: Path, dry_run: bool) -> dict[str, object]:
    before_hash = sha256(docx)
    before_members = package_members(docx)
    document = Document(docx)
    before_tables = table_hashes(document)
    matches = [paragraph for paragraph in document.paragraphs if visible_text(paragraph) == OLD]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one exact stale paragraph, found {len(matches)}")
    replace_text(matches[0], NEW)

    with tempfile.NamedTemporaryFile(prefix=f".{docx.name}.", suffix=".tmp.docx", dir=docx.parent, delete=False) as handle:
        staged = Path(handle.name)
    try:
        document.save(staged)
        verified = Document(staged)
        if sum(visible_text(paragraph) == NEW for paragraph in verified.paragraphs) != 1:
            raise RuntimeError("Replacement paragraph verification failed")
        if table_hashes(verified) != before_tables:
            raise RuntimeError("An Appendix table changed during paragraph-only update")
        after_members = package_members(staged)
        changed = sorted(
            name for name in set(before_members) | set(after_members)
            if before_members.get(name) != after_members.get(name)
        )
        forbidden = [name for name in changed if name.startswith("word/") and name != "word/document.xml"]
        if forbidden:
            raise RuntimeError(f"Protected Word members changed: {forbidden}")
        receipt: dict[str, object] = {
            "status": "dry-run" if dry_run else "applied",
            "sha256_before": before_hash,
            "sha256_after": sha256(staged),
            "changed_package_members": changed,
            "tables_unchanged": True,
            "backup": None,
        }
        if dry_run:
            return receipt
        backup_dir = docx.parent / ".kila-backups"
        backup_dir.mkdir(parents=True, exist_ok=True)
        stamp = dt.datetime.now().strftime("%Y%m%dT%H%M%S%f")
        backup = backup_dir / f"Appendix.before-r3c6-supplement.{stamp}.docx"
        shutil.copy2(docx, backup)
        if sha256(backup) != before_hash:
            raise RuntimeError("Backup hash mismatch")
        os.replace(staged, docx)
        receipt["backup"] = str(backup)
        return receipt
    finally:
        if staged.exists():
            staged.unlink()


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--docx", type=Path, default=Path("Rev/revision/Appendix.docx"))
    parser.add_argument("--apply", action="store_true")
    args = parser.parse_args()
    print(json.dumps(update(args.docx, dry_run=not args.apply), indent=2))


if __name__ == "__main__":
    main()
