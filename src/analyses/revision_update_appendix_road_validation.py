#!/usr/bin/env python3
"""Apply the approved Reviewer 2 Comment 7 updates to Appendix.docx.

The Appendix is a directly editable supplement rather than the tracked-change
main manuscript.  This script creates a timestamped backup, performs only the
approved B2/B5/B6b and explanatory-paragraph updates, validates the staged
DOCX, and atomically replaces the Appendix.
"""

from __future__ import annotations

from datetime import datetime, timezone
import hashlib
import os
from pathlib import Path
import shutil
import tempfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import pandas as pd


ROOT = Path(__file__).resolve().parents[2]
APPENDIX = ROOT / "Rev/revision/Appendix.docx"
BACKUP_DIR = ROOT / "Rev/revision/.kila-backups"
WORKBOOK = ROOT / "data/results/tables/Table_road_disruption_validation.xlsx"
TRANSFER = (
    ROOT
    / "data/exp/revision/reviewer-2-comment-7/transfer_event_clustered_validation.csv"
)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def paragraph_start(document: Document, prefix: str):
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text.startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph beginning {prefix!r}; found {len(matches)}.")
    return matches[0]


def set_paragraph(paragraph, text: str) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)


def set_cell(cell, value: object, *, bold: bool = False, size: float = 8.0) -> None:
    cell.text = "" if value is None else str(value)
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(size)
            run.bold = bold


def build_table(document: Document, data: pd.DataFrame, style, font_size: float = 8.0):
    table = document.add_table(rows=len(data) + 1, cols=len(data.columns))
    table.style = style
    table.autofit = True
    for column, label in enumerate(data.columns):
        set_cell(table.cell(0, column), label, bold=True, size=font_size)
    for row_index, values in enumerate(data.itertuples(index=False), start=1):
        for column, value in enumerate(values):
            set_cell(table.cell(row_index, column), value, size=font_size)
    return table


def format_correspondence(table: pd.DataFrame) -> pd.DataFrame:
    formatted = table.copy()
    formatted["Episode-Weighted Concordance"] = formatted[
        "Episode-Weighted Concordance"
    ].map(lambda value: f"{float(value):.3f}")
    formatted["Rank Correlation vs Heavy"] = formatted[
        "Rank Correlation vs Heavy"
    ].map(lambda value: "" if pd.isna(value) else f"{float(value):.3f}")
    compact_interpretation = {
        "Road score — Moderate rainfall": "Near-unity ranking; no spatial reprioritization.",
        "Road score — Heavy rainfall": "Supplementary dry-event correspondence only.",
        "Road score — Extreme rainfall": "Near-unity ranking; no spatial reprioritization.",
        "Road warning-zone exposure baseline": "Comparator using the same matched controls.",
        "Road-length baseline": "Comparator using the same matched controls.",
        "Heavy minus warning-zone baseline": "CI includes zero; no established superiority.",
        "Heavy minus road-length baseline": "CI includes zero; no established superiority.",
        "Heavy road score — Yatsushiro bound 0.70": "Analytical municipality-wide bound; not official.",
        "Heavy road score — Yatsushiro bound 0.80": "Analytical municipality-wide bound; not official.",
    }
    formatted["Permitted Interpretation"] = formatted["Specification"].map(
        compact_interpretation
    )
    return formatted


def compact_episode_table(table: pd.DataFrame) -> pd.DataFrame:
    compact = table[
        [
            "Municipality",
            "Process Reason",
            "Restriction Start Time",
            "Hours After Earthquake",
            "Repeated Snapshots",
            "Matched Sections",
            "Nearest JMA Station",
            "Nearest-Station Distance (km)",
            "Preceding 72 h Rainfall (mm)",
            "Trigger Classification",
        ]
    ].copy()
    compact["Restriction Start Time"] = (
        pd.to_datetime(compact["Restriction Start Time"], errors="coerce")
        .dt.strftime("%Y-%m-%d %H:%M")
    )
    compact["Hours After Earthquake"] = compact["Hours After Earthquake"].map(
        lambda value: f"{float(value):.2f}"
    )
    compact["Preceding 72 h Rainfall (mm)"] = compact[
        "Preceding 72 h Rainfall (mm)"
    ].map(lambda value: f"{float(value):.1f}")
    compact["Snapshots / sections"] = (
        compact["Repeated Snapshots"].astype(int).astype(str)
        + " / "
        + compact["Matched Sections"].astype(int).astype(str)
    )
    compact["JMA station; distance; 72 h rain"] = compact.apply(
        lambda row: (
            f"{row['Nearest JMA Station']}; "
            f"{float(row['Nearest-Station Distance (km)']):.1f} km; "
            f"{row['Preceding 72 h Rainfall (mm)']} mm"
        ),
        axis=1,
    )
    compact["Trigger Classification"] = compact["Trigger Classification"].replace(
        {
            "Direct-earthquake consistent; trigger not explicitly coded":
                "Direct-earthquake consistent; source-unconfirmed"
        }
    )
    return compact[
        [
            "Municipality",
            "Process Reason",
            "Restriction Start Time",
            "Hours After Earthquake",
            "Snapshots / sections",
            "JMA station; distance; 72 h rain",
            "Trigger Classification",
        ]
    ].rename(
        columns={
            "Process Reason": "Process",
            "Restriction Start Time": "Start (JST)",
            "Hours After Earthquake": "Hours after earthquake",
            "Trigger Classification": "Trigger classification",
        }
    )


def update_appendix(document: Document) -> None:
    correspondence = pd.read_excel(WORKBOOK, sheet_name="Correspondence", header=1)
    funnel = pd.read_excel(WORKBOOK, sheet_name="Evidence Funnel", header=1)
    episodes = pd.read_excel(WORKBOOK, sheet_name="Episode Audit", header=1)
    transfer = pd.read_csv(TRANSFER)
    if correspondence.shape != (9, 6) or funnel.shape != (11, 2) or episodes.shape != (10, 11):
        raise RuntimeError("Formal workbook dimensions do not match the approved Appendix update.")
    if transfer.shape[0] != 15:
        raise RuntimeError("Expected 15 transfer-sensitivity specifications.")

    road_text = paragraph_start(document, "Appendix Table B2 reports")
    set_paragraph(
        road_text,
        "Appendix Tables B2a–B2c separate the source-record funnel, physical restriction "
        "episodes, and matched road sections. Among 680 official snapshots, 175 carry "
        "rockfall, slope-collapse, or sediment-inflow process reasons; these resolve to 14 "
        "physical episodes before spatial matching and 10 retained episodes linked to 94 "
        "road sections, of which 93 have eligible controls. All ten episodes began "
        "0.55–24.38 h after the earthquake, contain no explicit rainfall-trigger term, and "
        "have 0 mm over the preceding 72 h across the ten-station JMA audit. They are "
        "therefore direct-earthquake consistent but source-unconfirmed and provide only "
        "supplementary terrain-to-road ranking correspondence. Heavy episode-weighted "
        "concordance is 0.723 (episode-cluster bootstrap 95% CI 0.598–0.840), compared "
        "with 0.698 (0.554–0.845) for road length and 0.545 (0.409–0.683) for warning-zone "
        "exposure; both paired contrast intervals include zero."
    )

    transfer_text = paragraph_start(document, "Across 15 specifications")
    set_paragraph(
        transfer_text,
        "Across 15 specifications, the strict joint boundary gives the lowest supported-road "
        "rank correlation (0.676) and top-1% overlap (0.428), while episode-weighted "
        "correspondence spans 0.711–0.741. The main instability arises from neighborhood "
        "reach and relief-based support rather than from the continuous distance, alignment, "
        "or relief weights alone. Because the correspondence sample contains only ten dry, "
        "earthquake-proximate physical episodes, this range is a supplementary consistency "
        "check and not rainfall-trigger validation. Table B7 propagates the joint boundaries "
        "using each setting's own Heavy 85th-percentile candidate set and 99.5th-percentile "
        "closure-mapping upper bound. Heavy expected isolated population spans "
        "523.5–2,256.3 residents around the central 1,121.7, so exact priorities and "
        "consequence magnitude remain conditional on the transfer specification."
    )

    b2_title = paragraph_start(document, "Appendix Table B2. Road-disruption validation")
    old_b2 = document.tables[2]
    table_style = old_b2.style
    set_paragraph(b2_title, "Appendix Table B2a. Road-restriction evidence funnel and trigger audit")
    funnel_table = build_table(document, funnel, table_style, font_size=8.5)
    old_b2._tbl.addprevious(funnel_table._tbl)
    old_b2._element.getparent().remove(old_b2._element)

    b2b_title = document.add_paragraph(
        "Appendix Table B2b. Event-weighted road-restriction correspondence"
    )
    b2b_title.style = b2_title.style
    b2b_table = build_table(
        document,
        format_correspondence(correspondence),
        table_style,
        font_size=7.0,
    )
    funnel_table._tbl.addnext(b2b_title._p)
    b2b_title._p.addnext(b2b_table._tbl)

    b2c_title = document.add_paragraph(
        "Appendix Table B2c. Restriction-episode trigger audit"
    )
    b2c_title.style = b2_title.style
    b2c_title.paragraph_format.page_break_before = True
    b2c_table = build_table(
        document,
        compact_episode_table(episodes),
        table_style,
        font_size=6.5,
    )
    b2b_table._tbl.addnext(b2c_title._p)
    b2c_title._p.addnext(b2c_table._tbl)

    b5 = next(
        table
        for table in document.tables
        if table.cell(0, 0).text.startswith("Table B5. Baseline-Threshold")
    )
    matched_rows = [
        row
        for row in b5.rows
        if "Matched road-evidence concordance" in row.cells[1].text
    ]
    if len(matched_rows) != 1:
        raise RuntimeError("Could not identify the unique Table B5 correspondence row.")
    b5_values = [
        "Rainfall parameters",
        "Episode-weighted road-restriction correspondence",
        "0.723",
        "0.711–0.726",
        "Range across 15 combinations",
        "10 physical episodes",
        "Supplementary dry-event correspondence; not rainfall-trigger validation",
    ]
    for cell, value in zip(matched_rows[0].cells, b5_values):
        set_cell(cell, value, size=7.5)

    b6b = next(
        table
        for table in document.tables
        if table.cell(0, 0).text.startswith("Table B6b. Road-ranking")
    )
    set_cell(
        b6b.cell(0, 0),
        "Table B6b. Road-ranking and episode-correspondence sensitivity",
        bold=True,
        size=8.0,
    )
    set_cell(
        b6b.cell(1, 5),
        "Episode-weighted concordance (95% CI)",
        bold=True,
        size=7.5,
    )
    transfer_lookup = transfer.set_index("Specification")
    ordered_keys = [
        "central",
        "radius_2",
        "radius_4",
        "relief_5",
        "relief_20",
        "alignment_0",
        "alignment_05",
        "decay_15",
        "decay_40",
        "relief_scale_50",
        "relief_scale_150",
        "midpoint",
        "five_points",
        "strict_joint",
        "permissive_joint",
    ]
    for row, key in zip(b6b.rows[2:], ordered_keys):
        values = transfer_lookup.loc[key]
        display = (
            f"{values['Episode-Weighted Concordance']:.3f} "
            f"({values['Episode-Cluster Bootstrap 95% CI Low']:.3f}–"
            f"{values['Episode-Cluster Bootstrap 95% CI High']:.3f})"
        )
        set_cell(row.cells[5], display, size=7.5)
    for row_index, row in enumerate(b6b.rows):
        for cell in row.cells:
            set_cell(
                cell,
                cell.text,
                bold=row_index in (0, 1),
                size=6.8 if row_index > 0 else 7.2,
            )


def verify(path: Path) -> None:
    document = Document(path)
    full_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    required = (
        "Appendix Table B2a. Road-restriction evidence funnel and trigger audit",
        "Appendix Table B2b. Event-weighted road-restriction correspondence",
        "Appendix Table B2c. Restriction-episode trigger audit",
        "cannot validate rainfall-triggered road disruption",
    )
    if not all(token in full_text for token in required[:3]):
        raise RuntimeError("One or more approved Appendix B2 titles are missing.")
    table_text = "\n".join(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    for token in (
        "Direct-earthquake-consistent episodes; trigger source-unconfirmed",
        "Heavy minus road-length baseline",
        "Episode-weighted road-restriction correspondence",
        "Episode-weighted concordance (95% CI)",
    ):
        if token not in table_text:
            raise RuntimeError(f"Missing approved Appendix content: {token}")
    if "93/94 / 870" in table_text or "Section-Bootstrap 95% CI" in table_text:
        raise RuntimeError("Superseded section-level validation content remains in Appendix tables.")
    if len(document.tables) != 14:
        raise RuntimeError(f"Expected 14 Appendix tables after B2 expansion; found {len(document.tables)}.")


def main() -> None:
    if not APPENDIX.exists():
        raise FileNotFoundError(APPENDIX)
    before_hash = sha256(APPENDIX)
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%S%f")
    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    backup = BACKUP_DIR / f"Appendix.{timestamp}.reviewer-2-comment-7.docx"
    shutil.copy2(APPENDIX, backup)

    document = Document(APPENDIX)
    update_appendix(document)
    with tempfile.NamedTemporaryFile(
        dir=APPENDIX.parent,
        suffix=".docx",
        delete=False,
    ) as handle:
        staged = Path(handle.name)
    try:
        document.save(staged)
        verify(staged)
        os.replace(staged, APPENDIX)
    finally:
        if staged.exists():
            staged.unlink()
    after_hash = sha256(APPENDIX)
    if before_hash == after_hash:
        raise RuntimeError("Appendix hash did not change after the approved update.")
    print(f"Appendix SHA-256 before: {before_hash}")
    print(f"Appendix SHA-256 after:  {after_hash}")
    try:
        backup_display = backup.relative_to(ROOT)
    except ValueError:
        backup_display = backup
    print(f"Backup: {backup_display}")
    print("Appendix verification: passed")


if __name__ == "__main__":
    main()
