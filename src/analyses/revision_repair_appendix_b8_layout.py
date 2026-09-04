#!/usr/bin/env python3
"""Repair only the approved Appendix Table B8 cell layout transactionally."""

from __future__ import annotations

import argparse
import datetime as dt
import hashlib
import json
import os
from pathlib import Path
import shutil
import tempfile
import zipfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


TITLE = "Table B8. Spatial closure-dependence sensitivity"


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def table_hashes(document: Document) -> list[str]:
    return [hashlib.sha256(t._tbl.xml.encode("utf-8")).hexdigest() for t in document.tables]


def members(path: Path) -> dict[str, bytes]:
    with zipfile.ZipFile(path) as archive:
        bad = archive.testzip()
        if bad:
            raise RuntimeError(f"Corrupt DOCX member: {bad}")
        return {name: archive.read(name) for name in archive.namelist()}


def apply(docx: Path, dry_run: bool) -> dict[str, object]:
    before_hash = sha256(docx)
    before_members = members(docx)
    document = Document(docx)
    tables = document.tables
    matches = [
        (index, table)
        for index, table in enumerate(tables)
        if table.cell(0, 0).text == TITLE
    ]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one Table B8, found {len(matches)}")
    target_index, target = matches[0]
    if (len(target.rows), len(target.columns)) != (17, 10):
        raise RuntimeError("Unexpected Table B8 shape")
    all_before = table_hashes(document)
    target_text_before = [[cell.text for cell in row.cells] for row in target.rows]

    for row_index, row in enumerate(target.rows):
        for column_index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            no_wrap = tc_pr.find(qn("w:noWrap"))
            if no_wrap is not None:
                tc_pr.remove(no_wrap)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.left_indent = Pt(0)
                paragraph.paragraph_format.right_indent = Pt(0)
                paragraph.paragraph_format.first_line_indent = Pt(0)
                if row_index >= 2:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    with tempfile.NamedTemporaryFile(
        prefix=f".{docx.name}.", suffix=".tmp.docx", dir=docx.parent, delete=False
    ) as handle:
        temp = Path(handle.name)
    try:
        document.save(temp)
        check = Document(temp)
        after_hashes = table_hashes(check)
        if len(after_hashes) != len(all_before):
            raise RuntimeError("Appendix table count changed")
        for index, (before, after) in enumerate(zip(all_before, after_hashes)):
            if index != target_index and before != after:
                raise RuntimeError(f"Protected Appendix table {index + 1} changed")
        check_target = check.tables[target_index]
        if [[cell.text for cell in row.cells] for row in check_target.rows] != target_text_before:
            raise RuntimeError("Table B8 text changed during layout repair")
        after_members = members(temp)
        changed = sorted(
            name for name in set(before_members) | set(after_members)
            if before_members.get(name) != after_members.get(name)
        )
        forbidden = [
            name for name in changed
            if name.startswith("word/") and name != "word/document.xml"
        ]
        if forbidden:
            raise RuntimeError(f"Protected package members changed: {forbidden}")
        receipt: dict[str, object] = {
            "status": "dry-run" if dry_run else "applied",
            "sha256_before": before_hash,
            "sha256_after": sha256(temp),
            "changed_package_members": changed,
            "table_b8_shape": [17, 10],
            "table_b8_text_unchanged": True,
            "other_tables_unchanged": True,
            "backup": None,
        }
        if dry_run:
            return receipt
        backup_dir = docx.parent / ".kila-backups"
        backup_dir.mkdir(parents=True, exist_ok=True)
        stamp = dt.datetime.now().strftime("%Y%m%dT%H%M%S%f")
        backup = backup_dir / f"{docx.stem}.{stamp}.reviewer-3-comment-5.part-06-layout.docx"
        shutil.copy2(docx, backup)
        if sha256(backup) != before_hash:
            raise RuntimeError("Appendix layout backup hash mismatch")
        os.replace(temp, docx)
        receipt["backup"] = str(backup)
        if sha256(docx) != receipt["sha256_after"]:
            raise RuntimeError("Final Appendix hash differs from validated staged file")
        return receipt
    finally:
        if temp.exists():
            temp.unlink()


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--docx", type=Path, required=True)
    parser.add_argument("--dry-run", action="store_true")
    args = parser.parse_args()
    print(json.dumps(apply(args.docx, args.dry_run), indent=2))


if __name__ == "__main__":
    main()
