#!/usr/bin/env python3
"""Insert the approved Reviewer 4 Comment 5 Appendix paragraph and Table B11."""

from __future__ import annotations

import argparse
import datetime as dt
import hashlib
import json
import os
from pathlib import Path
import shutil
import tempfile
import zipfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from openpyxl import load_workbook

from revision_update_appendix_r2c5 import (
    build_table,
    insert_paragraph_after,
    set_cell_text,
    set_row_cant_split,
)


HEADING = "Intervention Portfolio Robustness"
NARRATIVE = (
    "Appendix Table B11 treats the 26 unresolved announcements as a declared "
    "placement sensitivity rather than recovered data. In 50 replicates, the 18 "
    "Yatsushiro, six Uki, and two Hikawa records are distributed uniformly without "
    "replacement over eligible populated 125 m meshes in their source municipality "
    "and evaluated under five network seeds × 250 common Heavy-scenario closure draws. "
    "The expected affected-population range changes by −11.6% to −0.5% relative to the "
    "matched observed-only diagnostic, but community-frequency correlation and Top-30 "
    "overlap fall to 0.725 and 66.7% at their lower bounds. A same-seed test changes only "
    "water destinations and leaves shelter, fire-service, and municipal-facility loss, "
    "excess-time, and baseline arrays identical in 9/9 checks. This is a hypothetical "
    "sensitivity and not evidence of the actual missing locations."
)
EXPECTED_TITLE = "Appendix Table B11. Emergency-water missing-location sensitivity"
EXPECTED_HEADERS = [
    "Metric",
    "Observed-only reference",
    "Uniform P05",
    "Uniform median",
    "Uniform P95",
    "Interpretation",
]
WIDTHS = [1.35, 0.82, 0.70, 0.76, 0.70, 2.05]


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def package_members(path: Path) -> dict[str, bytes]:
    with zipfile.ZipFile(path) as archive:
        bad = archive.testzip()
        if bad:
            raise RuntimeError(f"Corrupt DOCX member: {bad}")
        return {name: archive.read(name) for name in archive.namelist()}


def table_hash(table) -> str:
    return hashlib.sha256(table._tbl.xml.encode("utf-8")).hexdigest()


def table_text(table) -> list[list[str]]:
    return [[cell.text for cell in row.cells] for row in table.rows]


def bookmark_fingerprint(document: Document) -> str:
    parts: list[str] = []
    for element in document.element.body.iter():
        if element.tag in {qn("w:bookmarkStart"), qn("w:bookmarkEnd")}:
            parts.append(f"{element.tag}:{sorted(element.attrib.items())}")
    return hashlib.sha256("\n".join(parts).encode("utf-8")).hexdigest()


def workbook_values(path: Path) -> tuple[str, list[str], list[list[str]], str]:
    workbook = load_workbook(path, data_only=True, read_only=True)
    sheet = workbook[workbook.sheetnames[0]]
    rows = list(sheet.iter_rows(values_only=True))
    title = str(rows[0][0])
    headers = [str(value) for value in rows[1][:6]]
    if title != EXPECTED_TITLE:
        raise RuntimeError(f"Unexpected Table B11 title: {title}")
    if headers != EXPECTED_HEADERS:
        raise RuntimeError(f"Unexpected Table B11 headers: {headers}")

    data_rows: list[list[str]] = []
    for row in rows[2:8]:
        metric = str(row[0])
        values: list[str] = [metric]
        for column, value in enumerate(row[1:5], start=1):
            if value is None:
                values.append("—")
            elif metric in {"Expected affected population", "Baseline-eligible population", "Newly eligible population"}:
                values.append(f"{float(value):,.1f}" if metric == "Expected affected population" else f"{float(value):,.0f}")
            elif metric == "Top-30 population-burden overlap":
                values.append(f"{100 * float(value):.1f}%")
            elif metric == "Community-frequency Spearman correlation":
                values.append(f"{float(value):.3f}")
            else:
                values.append(str(value))
        values.append(str(row[5]))
        data_rows.append(values)

    if len(data_rows) != 6 or data_rows[-1][0] != "Non-water class invariance":
        raise RuntimeError("Unexpected Table B11 data rows")
    note = str(rows[9][0])
    if "canonical five-seed × 1,000-draw estimate is 7,789.8" not in note:
        raise RuntimeError("Table B11 note lacks the canonical diagnostic")
    return title, headers, data_rows, note


def build_b11(document: Document, workbook_path: Path):
    title, headers, rows, note = workbook_values(workbook_path)
    table = build_table(document, title, headers, rows, WIDTHS, set())
    table.autofit = False
    for column, width in zip(table.columns, WIDTHS):
        column.width = Inches(width)
        for cell in column.cells:
            cell.width = Inches(width)

    note_row = table.add_row()
    note_cell = note_row.cells[0].merge(note_row.cells[-1])
    set_cell_text(
        note_cell,
        note,
        bold=False,
        size=6.0,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        color="404040",
        shading="F4F7F9",
    )
    set_row_cant_split(note_row)
    return table, [
        [title] * len(headers),
        headers,
        *rows,
        [note] * len(headers),
    ]


def update(docx: Path, workbook_path: Path, dry_run: bool) -> dict[str, object]:
    before_hash = sha256(docx)
    before_members = package_members(docx)
    document = Document(docx)
    original_table_hashes = [table_hash(table) for table in document.tables]
    original_bookmarks = bookmark_fingerprint(document)

    if any(paragraph.text == NARRATIVE for paragraph in document.paragraphs):
        raise RuntimeError("Reviewer 4 Comment 5 Appendix narrative is already present")
    if any(table.cell(0, 0).text == EXPECTED_TITLE for table in document.tables):
        raise RuntimeError("Appendix Table B11 already exists")
    heading_matches = [paragraph for paragraph in document.paragraphs if paragraph.text == HEADING]
    if len(heading_matches) != 1:
        raise RuntimeError(f"Expected one Appendix heading anchor, found {len(heading_matches)}")

    heading = heading_matches[0]
    narrative = insert_paragraph_after(heading._p, NARRATIVE, "Normal")
    heading._p.addprevious(narrative)
    table, expected_table_text = build_b11(document, workbook_path)
    heading._p.addprevious(table._tbl)
    heading._p.addprevious(OxmlElement("w:p"))

    with tempfile.NamedTemporaryFile(
        prefix=f".{docx.name}.", suffix=".tmp.docx", dir=docx.parent, delete=False
    ) as handle:
        staged = Path(handle.name)
    try:
        document.save(staged)
        verified = Document(staged)
        if sum(paragraph.text == NARRATIVE for paragraph in verified.paragraphs) != 1:
            raise RuntimeError("Saved Appendix narrative verification failed")
        matches = [
            table for table in verified.tables if table.cell(0, 0).text == EXPECTED_TITLE
        ]
        if len(matches) != 1:
            raise RuntimeError(f"Expected one saved Appendix Table B11, found {len(matches)}")
        checked = matches[0]
        if (len(checked.rows), len(checked.columns)) != (9, 6):
            raise RuntimeError("Appendix Table B11 shape verification failed")
        if table_text(checked) != expected_table_text:
            raise RuntimeError("Appendix Table B11 does not match the validated workbook")

        new_hash = table_hash(checked)
        unchanged_hashes = [
            value for value in [table_hash(table) for table in verified.tables] if value != new_hash
        ]
        if unchanged_hashes != original_table_hashes:
            raise RuntimeError("A pre-existing Appendix table changed")
        if bookmark_fingerprint(verified) != original_bookmarks:
            raise RuntimeError("Appendix bookmark fingerprint changed")

        after_members = package_members(staged)
        changed_members = sorted(
            name
            for name in set(before_members) | set(after_members)
            if before_members.get(name) != after_members.get(name)
        )
        unexpected = [
            name
            for name in changed_members
            if name not in {"word/document.xml", "docProps/core.xml"}
        ]
        if unexpected:
            raise RuntimeError(f"Protected Appendix package members changed: {unexpected}")

        receipt: dict[str, object] = {
            "status": "dry-run" if dry_run else "applied",
            "sha256_before": before_hash,
            "sha256_after": sha256(staged),
            "changed_package_members": changed_members,
            "paragraphs_inserted": 1,
            "table_b11_shape": [9, 6],
            "preexisting_table_count": len(original_table_hashes),
            "preexisting_tables_unchanged": True,
            "bookmarks_preserved": True,
            "backup": None,
        }
        if dry_run:
            return receipt

        backup_dir = docx.parent / ".kila-backups"
        backup_dir.mkdir(parents=True, exist_ok=True)
        stamp = dt.datetime.now().strftime("%Y%m%dT%H%M%S%f")
        backup = backup_dir / f"Appendix.before-r4c5.{stamp}.docx"
        shutil.copy2(docx, backup)
        if sha256(backup) != before_hash:
            raise RuntimeError("Appendix backup hash mismatch")
        os.replace(staged, docx)
        receipt["backup"] = str(backup)
        if sha256(docx) != receipt["sha256_after"]:
            raise RuntimeError("Final Appendix hash differs from validated staged file")
        return receipt
    finally:
        if staged.exists():
            staged.unlink()


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--docx", type=Path, default=Path("Rev/revision/Appendix.docx"))
    parser.add_argument(
        "--workbook",
        type=Path,
        default=Path(
            "data/exp/revision/reviewer-4-comment-5/"
            "Table_emergency_water_missing_location_sensitivity.xlsx"
        ),
    )
    parser.add_argument("--apply", action="store_true")
    args = parser.parse_args()
    print(json.dumps(update(args.docx, args.workbook, dry_run=not args.apply), indent=2))


if __name__ == "__main__":
    main()
