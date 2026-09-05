"""Insert one approved B14/B15 table at a time; preserve all old Appendix XML."""
from __future__ import annotations
import argparse
import datetime as dt
import json
import os
from pathlib import Path
import shutil
import tempfile

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import load_workbook

from revision_build_terrain_sensitivity_tables import approved_table, ROOT, OUT, TABLES, STEMS, PROPOSAL
from revision_update_appendix_r2c5 import build_table, page_break_paragraph
from revision_update_appendix_r2c6 import sha256, package_members, bookmark_fingerprint


def update(number, apply):
    docx = ROOT/'Rev/revision/Appendix.docx'
    title, headers, rows, note = approved_table(number)
    workbook_path = TABLES/(STEMS[number]+'.xlsx')
    workbook = load_workbook(workbook_path, data_only=True)
    sheet = workbook.active
    assert sheet['A1'].value == title
    assert [sheet.cell(2,c).value for c in range(1,len(headers)+1)] == headers
    if number == 15:
        assert [[sheet.cell(r,c).value for c in range(1,5)] for r in range(3,11)] == rows
    else:
        for r,expected in enumerate(rows,3):
            actual=[sheet.cell(r,1).value,*[f'{sheet.cell(r,c).value:,.1f}' for c in (2,3)]]
            assert actual == expected[:3]
            assert f'{sheet.cell(r,4).value:.1f}' == expected[3].split(' [')[0]
            assert [f'{sheet.cell(r,c).value:.1f}' for c in (5,6,7)] == expected[4:]
    before_hash = sha256(docx)
    members_before = package_members(docx)
    document = Document(docx)
    original_elements = [element.xml for element in document.element.body]
    bookmarks_before = bookmark_fingerprint(document)
    assert not any(table.cell(0,0).text == title for table in document.tables)
    c1 = [p for p in document.paragraphs if p.text == 'Appendix Table C1. Municipality isolation and service-loss summary']
    assert len(c1) == 1
    if number == 14:
        anchor = PROPOSAL.read_text().split('Exact anchor paragraph: "',1)[1].split('"',1)[0]
    else:
        anchor = approved_table(14)[3]
    anchors = [p for p in document.paragraphs if p.text == anchor]
    assert len(anchors) == 1
    # Before C1, only previously added page-break paragraphs may follow the anchor.
    sibling = anchors[0]._p.getnext()
    while sibling is not c1[0]._p:
        assert sibling is not None and not sibling.xpath('.//w:t')
        sibling = sibling.getnext()
    widths = [1.60,1.05,.70,1.50,1.15,1.20,1.30] if number==14 else [1.40,2.05,1.85,3.20]
    table = build_table(document,title,headers,rows,widths,set())
    for index,row in enumerate(table.rows):
        for col,cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before=Pt(3)
                paragraph.paragraph_format.space_after=Pt(3)
                paragraph.paragraph_format.line_spacing=1.0
                if number==15 and index>1:
                    paragraph.alignment=WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size=Pt(10 if index==0 else 8.5)
    note_paragraph = document.add_paragraph(note)
    note_paragraph.paragraph_format.space_before=Pt(8)
    note_paragraph.paragraph_format.space_after=Pt(0)
    note_paragraph.paragraph_format.line_spacing=1.0
    for run in note_paragraph.runs:
        run.font.size=Pt(9)
        run.italic=True
    cursor=anchors[0]._p
    new_nodes=[page_break_paragraph(), table._tbl, note_paragraph._p]
    # C1 already has pageBreakBefore; an extra break would create a blank page.
    for node in new_nodes:
        cursor.addnext(node)
        cursor=node
    body=document.element.body
    assert [element.xml for element in body if element not in new_nodes]==original_elements
    staged=None
    try:
        with tempfile.NamedTemporaryFile(prefix='.Appendix-r3c2-',suffix='.docx',dir=docx.parent,delete=False) as f:
            staged=Path(f.name)
        document.save(staged)
        verified=Document(staged)
        checked=[t for t in verified.tables if t.cell(0,0).text==title]
        assert len(checked)==1
        assert [[c.text for c in r.cells] for r in checked[0].rows] == [[title]*len(headers),headers,*rows]
        assert sum(p.text==note for p in verified.paragraphs)==1
        assert bookmark_fingerprint(verified)==bookmarks_before
        remaining=[]
        new_indices=[list(body).index(node) for node in new_nodes]
        for index,element in enumerate(verified.element.body):
            if index not in new_indices:
                remaining.append(element.xml)
        assert remaining==original_elements, 'A pre-existing Appendix body element changed'
        after_members=package_members(staged)
        changed=sorted(k for k in set(members_before)|set(after_members) if members_before.get(k)!=after_members.get(k))
        assert set(changed)<= {'word/document.xml','docProps/core.xml'},changed
        receipt=dict(status='applied' if apply else 'dry-run',part=f'part-{number-9:02d}',table=number,
                     sha256_before=before_hash,sha256_after=sha256(staged),
                     changed_members=changed,all_old_body_elements_unchanged=True,
                     bookmarks_preserved=True,workbook_sha256=sha256(workbook_path),backup=None)
        if apply:
            log=ROOT/'Rev/docs/revisionchanges.md'
            old_log=log.read_text()
            comment_section=old_log.split('## reviewer-3/comment-2',1)[1].split('\n## ',1)[0]
            assert f'### {receipt["part"]}' not in comment_section
            stamp=dt.datetime.now().strftime('%Y%m%dT%H%M%S%f')
            backup=docx.parent/'.kila-backups'/f'Appendix.before-r3c2-b{number}.{stamp}.docx'
            shutil.copy2(docx,backup)
            assert sha256(backup)==before_hash
            receipt['backup']=str(backup)
            os.replace(staged,docx)
            record=f'\n### {receipt["part"]}\n\n- Location: Appendix, after the preceding B-table notes and before C1\n- Mode: `appendix-table-add` (explicit direct-edit authorization)\n- Kila decisions: KILA-D-20260905-013; KILA-D-20260905-014\n- Timestamp: {dt.datetime.now().astimezone().isoformat()}\n- Before: no B{number} block; all existing content retained\n- Appendix SHA-256 before: `{before_hash}`\n- Appendix SHA-256 after: `{receipt["sha256_after"]}`\n- Backup: `{backup}`\n- Original body XML, tables, bookmarks and protected package members preserved: true\n- Workbook: `{workbook_path.relative_to(ROOT)}`\n- Workbook SHA-256: `{receipt["workbook_sha256"]}`\n- Exact added content:\n\n~~~~text\n{title}\n'+ '\n'.join(' | '.join(row) for row in [headers,*rows])+f'\n{note}\n~~~~\n'
            # The selected comment is the final heading; preserve the entire old log.
            assert not old_log.split('## reviewer-3/comment-2',1)[1].count('\n## ')
            log.write_text(old_log+record)
            (OUT/f'appendix-part-{number-9:02d}-receipt.json').write_text(json.dumps(receipt,indent=2)+'\n')
        return receipt
    finally:
        if staged and staged.exists():
            staged.unlink()


def repair_added_break():
    """Remove only our extra empty break before C1, keeping every old element."""
    docx=ROOT/'Rev/revision/Appendix.docx'
    before_hash=sha256(docx)
    assert before_hash=='e80a31d978b3a046f0e94600d7166cfac3fa955a296d4926d7e641851b311128'
    members_before=package_members(docx)
    document=Document(docx)
    c1=next(p for p in document.paragraphs if p.text.startswith('Appendix Table C1.'))
    node=c1._p.getprevious()
    assert not node.xpath('.//w:t') and len(node.xpath('./w:pPr/w:pageBreakBefore'))==1
    assert ''.join(node.getprevious().xpath('.//w:t/text()'))==approved_table(15)[3]
    expected=[element.xml for element in document.element.body if element is not node]
    document.element.body.remove(node)
    with tempfile.NamedTemporaryFile(prefix='.Appendix-r3c2-layout-',suffix='.docx',dir=docx.parent,delete=False) as f:
        staged=Path(f.name)
    try:
        document.save(staged)
        assert [element.xml for element in Document(staged).element.body]==expected
        after_members=package_members(staged)
        assert [k for k in members_before if members_before[k]!=after_members[k]]==['word/document.xml']
        stamp=dt.datetime.now().strftime('%Y%m%dT%H%M%S%f')
        backup=docx.parent/'.kila-backups'/f'Appendix.before-r3c2-layout.{stamp}.docx'
        shutil.copy2(docx,backup)
        assert sha256(backup)==before_hash
        after_hash=sha256(staged)
        os.replace(staged,docx)
        log=ROOT/'Rev/docs/revisionchanges.md'
        log.write_text(log.read_text()+f'\n### part-06-layout\n\n- Location: Appendix, new B15/C1 boundary\n- Mode: layout-only correction of agent-added blank break\n- Kila decisions: KILA-D-20260905-013; KILA-D-20260905-014\n- Before: one agent-added empty page-break paragraph immediately before C1, whose original pageBreakBefore already starts a page\n- After: remove only that new empty break; all words, tables and pre-existing XML unchanged\n- Appendix SHA-256 before: `{before_hash}`\n- Appendix SHA-256 after: `{after_hash}`\n- Backup: `{backup}`\n- Verification: every remaining body element byte-identical; all other package members byte-identical; rerender required\n')
        receipt=dict(before=before_hash,after=after_hash,backup=str(backup),only_added_blank_break_removed=True)
        (OUT/'appendix-layout-receipt.json').write_text(json.dumps(receipt,indent=2)+'\n')
        return receipt
    finally:
        if staged.exists(): staged.unlink()


if __name__=='__main__':
    parser=argparse.ArgumentParser()
    parser.add_argument('--table',type=int,choices=[14,15])
    parser.add_argument('--apply',action='store_true')
    parser.add_argument('--repair-added-break',action='store_true')
    args=parser.parse_args()
    assert args.repair_added_break or args.table
    print(json.dumps(repair_added_break() if args.repair_added_break else update(args.table,args.apply),indent=2))
