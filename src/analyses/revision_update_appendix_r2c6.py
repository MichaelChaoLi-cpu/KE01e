#!/usr/bin/env python3
"""Apply the approved Reviewer 2 Comment 6 Appendix update and Table B13."""

from __future__ import annotations

import argparse
import datetime as dt
import hashlib
import json
import os
from pathlib import Path
import shutil
import tempfile
import zipfile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from openpyxl import load_workbook

from revision_update_appendix_r2c5 import build_table, page_break_paragraph


ANCHOR = (
    "Appendix Table B3 shows that protected population increases with budget under "
    "Conservative, Central, and Optimistic assumptions. At the maximum budget of "
    "269.13 relative planning units, benefits are 31.2, 62.3, and 99.6 residents, "
    "while selection overlap with the Central portfolio is 84.0% in the Conservative "
    "setting and 80.6% in the Optimistic setting. Appendix Table B4 shows that the "
    "Central assigned-action and equal-cost consequence rankings both protect 62.3 "
    "residents, whereas hazard-only, emergency-route-only, and road-class-only rules "
    "protect 0.5, 0.3, and 0.0 residents, respectively. The comparison supports "
    "consequence-aware screening but not superiority over the equivalent consequence "
    "benchmark."
)
NARRATIVE = (
    "Appendix Table B13 separates intervention effectiveness, cost structure, and the "
    "Equation 17 attachment coefficient one family at a time under Heavy rainfall, the "
    "Primary Emergency Road target, a fixed 269.131-unit budget, five prespecified seeds, "
    "and 1,000 common-random-number draws per seed. Effectiveness assumptions preserve "
    "near-identical road ordering but change protected population from 33.3 to 90.7 "
    "residents around the Central 62.3. Cost alternatives yield 52.9–67.8 protected "
    "residents and reduce Top-30 overlap to 40.0% when fixed cost components are removed. "
    "Attachment-coefficient values from 0 to 0.50 retain the Central Top 30 and yield "
    "59.8–62.3 protected residents; the coefficient-zero all-road correlation is "
    "tie-sensitive because only 216 roads retain a positive score. The results therefore "
    "bound planning sensitivity but do not validate the numerical assumptions as local "
    "engineering performance or cost estimates."
)
EXPECTED_TITLE = (
    "Table B13. One-family-at-a-time sensitivity of intervention screening assumptions"
)
EXPECTED_HEADERS = [
    "Parameter family",
    "Setting",
    "Attachment coefficient (λ)",
    "Roads with positive score",
    "Score Spearman vs Central",
    "Top-30 overlap vs Central",
    "Selected roads",
    "Portfolio overlap vs Central",
    "Realized planning cost",
    "Action mix (R/C/A)",
    "Protected population, mean [seed range]",
    "Change vs Central",
]
DISPLAY_HEADERS = [
    "Parameter\nfamily",
    "Setting",
    "Attachment\ncoefficient (λ)",
    "Roads with\npositive score",
    "Score Spearman\nvs Central",
    "Top-30 overlap\nvs Central",
    "Selected\nroads",
    "Portfolio overlap\nvs Central",
    "Realized\nplanning cost",
    "Action mix\n(R/C/A)",
    "Protected population,\nmean [seed range]",
    "Change vs\nCentral",
]
WIDTHS = [0.80, 1.05, 0.55, 0.60, 0.65, 0.65, 0.55, 0.65, 0.65, 0.75, 1.05, 0.55]


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def package_members(path: Path) -> dict[str, bytes]:
    with zipfile.ZipFile(path) as archive:
        bad = archive.testzip()
        if bad:
            raise RuntimeError(f"Corrupt DOCX member: {bad}")
        return {name: archive.read(name) for name in archive.namelist()}


def table_hash(table) -> str:
    return hashlib.sha256(table._tbl.xml.encode("utf-8")).hexdigest()


def table_text(table) -> list[list[str]]:
    return [[cell.text for cell in row.cells] for row in table.rows]


def bookmark_fingerprint(document: Document) -> str:
    parts: list[str] = []
    for element in document.element.body.iter():
        if element.tag in {qn("w:bookmarkStart"), qn("w:bookmarkEnd")}:
            parts.append(f"{element.tag}:{sorted(element.attrib.items())}")
    return hashlib.sha256("\n".join(parts).encode("utf-8")).hexdigest()


def workbook_values(path: Path) -> tuple[list[list[str]], str]:
    workbook = load_workbook(path, data_only=True, read_only=True)
    sheet = workbook["Intervention sensitivity"]
    rows = list(sheet.iter_rows(values_only=True))
    title = str(rows[0][0])
    headers = [str(value) for value in rows[1][:12]]
    if title != EXPECTED_TITLE:
        raise RuntimeError(f"Unexpected Table B13 title: {title}")
    if headers != EXPECTED_HEADERS:
        raise RuntimeError(f"Unexpected Table B13 headers: {headers}")

    output: list[list[str]] = []
    for row in rows[2:13]:
        output.append(
            [
                str(row[0]),
                str(row[1]),
                f"{float(row[2]):.3f}",
                f"{int(row[3]):,}",
                f"{float(row[4]):.3f}",
                f"{100 * float(row[5]):.1f}%",
                f"{int(row[6]):,}",
                f"{100 * float(row[7]):.1f}%",
                f"{float(row[8]):.1f}",
                str(row[9]),
                str(row[10]),
                f"{100 * float(row[11]):+.1f}%" if float(row[11]) else "0.0%",
            ]
        )
    if len(output) != 11:
        raise RuntimeError(f"Expected 11 Table B13 data rows, found {len(output)}")
    note = str(rows[14][0])
    if not note.startswith("Notes: Heavy rainfall"):
        raise RuntimeError("Unexpected Table B13 note")
    return output, note


def build_b13(document: Document, workbook_path: Path):
    rows, note = workbook_values(workbook_path)
    table = build_table(
        document,
        EXPECTED_TITLE,
        DISPLAY_HEADERS,
        rows,
        WIDTHS,
        {2, 3, 4, 5, 6, 7, 8, 9, 11},
    )
    table.autofit = False
    for column, width in zip(table.columns, WIDTHS, strict=True):
        column.width = Inches(width)
        for cell in column.cells:
            cell.width = Inches(width)
    for cell in table.rows[1].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(4.6)
    for row in table.rows[2:]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(5.0)
    expected = [[EXPECTED_TITLE] * 12, DISPLAY_HEADERS, *rows]
    return table, expected, note


def insert_paragraph_after(previous, text: str, style: str):
    paragraph = Document().add_paragraph(text, style=style)
    element = paragraph._p
    previous.addnext(element)
    return element


def update(docx: Path, workbook_path: Path, dry_run: bool) -> dict[str, object]:
    before_hash = sha256(docx)
    before_members = package_members(docx)
    document = Document(docx)
    original_table_hashes = [table_hash(table) for table in document.tables]
    original_bookmarks = bookmark_fingerprint(document)

    matches = [paragraph for paragraph in document.paragraphs if paragraph.text == ANCHOR]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one exact Appendix anchor, found {len(matches)}")
    if any(table.cell(0, 0).text == EXPECTED_TITLE for table in document.tables):
        raise RuntimeError("Appendix Table B13 already exists")

    table, expected_table_text, note = build_b13(document, workbook_path)
    cursor = matches[0]._p
    insert_paragraph_after(cursor, NARRATIVE, "Normal")
    c1_titles = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text
        == "Appendix Table C1. Municipality isolation and service-loss summary"
    ]
    if len(c1_titles) != 1:
        raise RuntimeError(f"Expected one Appendix Table C1 title, found {len(c1_titles)}")
    c1_title = c1_titles[0]._p
    c1_title.addprevious(page_break_paragraph())
    c1_title.addprevious(table._tbl)
    note_paragraph = Document().add_paragraph(note, style="Normal")._p
    c1_title.addprevious(note_paragraph)
    for run in note_paragraph.xpath(".//w:r"):
        run_properties = run.find(qn("w:rPr"))
        if run_properties is None:
            from docx.oxml import OxmlElement

            run_properties = OxmlElement("w:rPr")
            run.insert(0, run_properties)
        from docx.oxml import OxmlElement

        italic = OxmlElement("w:i")
        run_properties.append(italic)

    with tempfile.NamedTemporaryFile(
        prefix=f".{docx.name}.", suffix=".tmp.docx", dir=docx.parent, delete=False
    ) as handle:
        staged = Path(handle.name)
    try:
        document.save(staged)
        verified = Document(staged)
        if sum(paragraph.text == NARRATIVE for paragraph in verified.paragraphs) != 1:
            raise RuntimeError("Saved Appendix narrative verification failed")
        if sum(paragraph.text == note for paragraph in verified.paragraphs) != 1:
            raise RuntimeError("Saved Appendix note verification failed")
        table_matches = [
            table for table in verified.tables if table.cell(0, 0).text == EXPECTED_TITLE
        ]
        if len(table_matches) != 1:
            raise RuntimeError(f"Expected one saved Appendix Table B13, found {len(table_matches)}")
        checked = table_matches[0]
        if (len(checked.rows), len(checked.columns)) != (13, 12):
            raise RuntimeError("Appendix Table B13 shape verification failed")
        if table_text(checked) != expected_table_text:
            raise RuntimeError("Appendix Table B13 does not match the validated workbook")

        new_hash = table_hash(checked)
        unchanged_hashes = [
            value for value in [table_hash(table) for table in verified.tables] if value != new_hash
        ]
        if unchanged_hashes != original_table_hashes:
            raise RuntimeError("A pre-existing Appendix table changed")
        if bookmark_fingerprint(verified) != original_bookmarks:
            raise RuntimeError("Appendix bookmark fingerprint changed")

        after_members = package_members(staged)
        changed_members = sorted(
            name
            for name in set(before_members) | set(after_members)
            if before_members.get(name) != after_members.get(name)
        )
        unexpected = [
            name for name in changed_members if name not in {"word/document.xml", "docProps/core.xml"}
        ]
        if unexpected:
            raise RuntimeError(f"Protected Appendix package members changed: {unexpected}")

        receipt: dict[str, object] = {
            "status": "dry-run" if dry_run else "applied",
            "sha256_before": before_hash,
            "sha256_after": sha256(staged),
            "changed_package_members": changed_members,
            "narratives_inserted": 1,
            "table_b13_shape": [13, 12],
            "preexisting_table_count": len(original_table_hashes),
            "preexisting_tables_unchanged": True,
            "bookmarks_preserved": True,
            "backup": None,
        }
        if dry_run:
            return receipt

        backup_dir = docx.parent / ".kila-backups"
        backup_dir.mkdir(parents=True, exist_ok=True)
        stamp = dt.datetime.now().strftime("%Y%m%dT%H%M%S%f")
        backup = backup_dir / f"Appendix.before-r2c6.{stamp}.docx"
        shutil.copy2(docx, backup)
        if sha256(backup) != before_hash:
            raise RuntimeError("Appendix backup hash mismatch")
        os.replace(staged, docx)
        receipt["backup"] = str(backup)
        if sha256(docx) != receipt["sha256_after"]:
            raise RuntimeError("Final Appendix hash differs from validated staged file")
        return receipt
    finally:
        if staged.exists():
            staged.unlink()


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--docx", type=Path, default=Path("Rev/revision/Appendix.docx"))
    parser.add_argument(
        "--workbook",
        type=Path,
        default=Path(
            "data/exp/revision/reviewer-2-comment-6/"
            "Table_intervention_parameter_sensitivity.xlsx"
        ),
    )
    parser.add_argument("--apply", action="store_true")
    args = parser.parse_args()
    print(json.dumps(update(args.docx, args.workbook, dry_run=not args.apply), indent=2))


if __name__ == "__main__":
    main()
