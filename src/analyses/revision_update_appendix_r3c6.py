#!/usr/bin/env python3
"""Apply the approved Reviewer 3 Comment 6 Appendix update transactionally."""

from __future__ import annotations

import argparse
from copy import deepcopy
import datetime as dt
import hashlib
import json
import os
from pathlib import Path
import shutil
import tempfile
import zipfile

import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import load_workbook

from revision_update_appendix_r2c5 import build_table, page_break_paragraph


PARAGRAPH_REPLACEMENTS = {
    "Baseline reconciliation is applied before network outcomes are calculated. Populated meshes must attach to an eligible road component, community population must reconcile to the retained census support, and the community must connect to the declared external target. Service pairs must also have a valid destination attachment and a baseline route. Disconnected communities, unresolved destinations, and baseline-unreachable service pairs remain explicitly non-evaluable rather than entering disruption totals.":
        "Baseline reconciliation is applied before network outcomes are calculated. Populated meshes must attach to an eligible road component, community population must reconcile to the retained census support, and the community must connect to the Primary Emergency Road backbone. Service pairs must also have a valid destination attachment and a baseline route. Disconnected communities, unresolved destinations, and baseline-unreachable service pairs remain explicitly non-evaluable rather than entering disruption totals.",
    "Across 15 specifications, the strict joint boundary gives the lowest supported-road rank correlation (0.676) and top-1% overlap (0.428), while episode-weighted correspondence spans 0.711–0.741. The main instability arises from neighborhood reach and relief-based support rather than from the continuous distance, alignment, or relief weights alone. Because the correspondence sample contains only ten dry, earthquake-proximate physical episodes, this range is a supplementary consistency check and not rainfall-trigger validation. Table B7 propagates the joint boundaries using each setting's own Heavy 85th-percentile candidate set and 99.5th-percentile closure-mapping upper bound. Heavy expected isolated population spans 523.5–2,256.3 residents around the central 1,121.7, so exact priorities and consequence magnitude remain conditional on the transfer specification.":
        "Across 15 specifications, the strict joint boundary gives the lowest supported-road rank correlation (0.676) and top-1% overlap (0.428), while episode-weighted correspondence spans 0.711–0.741. The main instability arises from neighborhood reach and relief-based support rather than from the continuous distance, alignment, or relief weights alone. Because the correspondence sample contains only ten dry, earthquake-proximate physical episodes, this range is a supplementary consistency check and not rainfall-trigger validation. Table B7 propagates the joint boundaries using each setting's own Heavy 85th-percentile candidate set and 99.5th-percentile closure-mapping upper bound. Heavy expected disconnected population spans 518.3–1,999.4 residents around the central 1,063.6, so exact priorities and consequence magnitude remain conditional on the transfer specification.",
    "Appendix Table B5 compares the official threshold geography with an all-area f = 1.00 baseline that switches off threshold retention while holding other inputs fixed. Expected isolated population increased from 85.9 to 189.1, 571.0 to 1,121.7, and 4,121.7 to 5,032.2 persons under Moderate, Heavy, and Extreme rainfall, respectively. Under Heavy rainfall, expected reachability losses increased from 272.8 to 609.9 persons for shelters, from 596.2 to 1,286.7 for fire services, and from 497.1 to 937.3 for municipal facilities. The emergency-water comparison remains conditional on the 10 geocoded source features.":
        "Appendix Table B5 compares the official threshold geography with an all-area (f=1.00) baseline that switches off threshold retention while holding other inputs fixed. Expected disconnected population increased from 85.9 to 189.1, 564.3 to 1,063.6, and 3,455.4 to 4,217.8 persons under Moderate, Heavy, and Extreme rainfall, respectively. Under Heavy rainfall, expected reachability losses increased from 272.8 to 609.9 persons for shelters, from 596.2 to 1,286.7 for fire services, and from 497.1 to 937.3 for municipal facilities. The emergency-water comparison remains conditional on the 10 geocoded source features.",
    "Road rankings remained similar between the all-area baseline and official geography under Heavy rainfall (Spearman correlation, 0.995; top-1% overlap, 0.879). Across five window-weight schemes and three γ values, the minimum road-ranking correlation was 0.989, minimum top-1% overlap was 0.879, and matched road-evidence concordance ranged from 0.632 to 0.652. Heavy expected isolated population ranged from 786.2 to 2,135.9 persons, or 0.701–1.904 times the central estimate of 1,121.7. Thus, location screening is comparatively robust to the tested rainfall-window and γ choices, whereas consequence magnitude remains sensitive to the rainfall parameterization; Appendix Tables B6 and B7 separately show that the slope-to-road influence-set boundaries materially affect exact road priorities and downstream magnitude.":
        "Road rankings remained similar between the all-area baseline and official geography under Heavy rainfall (Spearman correlation, 0.995; top-1% overlap, 0.879). Across five window-weight schemes and three γ values, the minimum road-ranking correlation was 0.989, minimum top-1% overlap was 0.879, and matched road-evidence concordance ranged from 0.632 to 0.652. Heavy expected disconnected population ranged from 765.7 to 1,874.6 persons, or 0.720–1.763 times the central estimate of 1,063.6. Thus, location screening is comparatively robust to the tested rainfall-window and γ choices, whereas consequence magnitude remains sensitive to the rainfall parameterization; Appendix Tables B6 and B7 separately show that the slope-to-road influence-set boundaries materially affect exact road priorities and downstream magnitude.",
    "Simulation-size and network-definition checks bound the isolation results. Changing the number of draws from 500 to 2,000 with one common seed produces a 0.006 difference in the 95th-percentile community isolation frequency. Alternative external-road targets place Heavy-scenario expected isolation between 1,044 and 1,108 residents, compared with the five-seed Central mean of 1,106.9 under the primary target. Municipality-wide Yatsushiro assignments of 0.70 and 0.80 bound the result between 1,041 and 1,193, while alternative closure mappings produce the wider range of 343–2,309 residents. The closure mapping is therefore a larger source of uncertainty than Monte Carlo convergence or the tested gateway definitions.":
        "Simulation-size and network-target checks bound the disconnection results. Changing the number of draws from 500 to 2,000 with one common seed produces a 0.006 difference in the 95th-percentile community isolation frequency. The Primary Emergency Road backbone contains 2,562 target roots and yields 1,063.6 expected disconnected residents under Heavy rainfall. Expanding the target to 2,977 Primary-plus-Secondary Emergency Road roots yields 992.7 residents, with community-frequency Spearman correlation 0.964 and top-30 burden overlap 90.0% relative to the primary definition. The former coast-inclusive boundary proxy is reported only as an audit comparator and reproduces the prior 1,121.7-resident result. Municipality-wide Yatsushiro assignments of 0.70 and 0.80 bound the revised primary result between 1,016.6 and 1,118.7, while alternative closure mappings produce the wider rounded range of 343–2,057 residents. Closure mapping is therefore a larger source of magnitude uncertainty than Monte Carlo convergence or the tested emergency-road target definition.",
    "Appendix Table B8 isolates spatial closure dependence while retaining each candidate section's central marginal closure propensity. The independent implementation exactly reproduces the existing five-seed simulation. Under the broad strong setting, mean expected isolated population changes by +9.2%, +15.7%, and −10.5% for Moderate, Heavy, and Extreme rainfall, respectively, and the corresponding per-draw 95th percentiles change by +8.7%, +15.8%, and −5.3%. Community-frequency rank correlations remain high, but lower top-30 overlap under Heavy rainfall shows that local preparedness priorities are more dependence-sensitive than the broad geographic ordering. Because the scale and correlation settings are not calibrated, the table reports sensitivity bounds rather than alternative forecasts.":
        "Appendix Table B8 isolates spatial closure dependence while retaining each candidate section's central marginal closure propensity. The independent implementation exactly reproduces the revised five-seed simulation. Under the broad strong setting, mean expected disconnected population changes by +9.2%, +16.0%, and −6.5% for Moderate, Heavy, and Extreme rainfall, respectively, and the corresponding per-draw 95th percentiles change by +8.7%, +18.0%, and −2.8%. Community-frequency rank correlations remain high, but 70.0% top-30 overlap under Heavy rainfall shows that local preparedness priorities are more dependence-sensitive than the broad geographic ordering. Because the scale and correlation settings are not calibrated, the table reports sensitivity bounds rather than alternative forecasts.",
    "Appendix Table C3 lists the top 30 communities using English municipality or ward names and geographic centroids. It retains mesh count, total and older population, candidate gateway-section count, isolation frequency under each rainfall scenario, Heavy-scenario expected isolated population, and the principal service loss. The entries are population clusters defined by the baseline road network, not administrative settlements or confirmed isolated locations.":
        "Appendix Table C3 lists the top 30 communities using English municipality or ward names and geographic centroids. It retains mesh count, total and older population, candidate connection-section count, isolation frequency under each rainfall scenario, Heavy-scenario expected disconnected population, and the principal service loss. The entries are population clusters defined by the baseline road network, not administrative settlements or confirmed isolated locations.",
}

OLD_A1_ROLE = "External-road targets and intervention stratification"
NEW_A1_ROLE = "Emergency-road-backbone targets and intervention stratification"


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def package_members(path: Path) -> dict[str, bytes]:
    with zipfile.ZipFile(path) as archive:
        bad = archive.testzip()
        if bad:
            raise RuntimeError(f"Corrupt DOCX member: {bad}")
        return {name: archive.read(name) for name in archive.namelist()}


def table_hash(table) -> str:
    return hashlib.sha256(table._tbl.xml.encode("utf-8")).hexdigest()


def paragraph_visible_text(paragraph) -> str:
    visible = {qn("w:t"), qn("w:delText"), qn("m:t")}
    return "".join(element.text or "" for element in paragraph._p.iter() if element.tag in visible)


def replace_paragraph_text(paragraph, new_text: str) -> None:
    p = paragraph._p
    p_pr = p.find(qn("w:pPr"))
    first_r_pr = None
    first_run = p.find(qn("w:r"))
    if first_run is not None:
        source_r_pr = first_run.find(qn("w:rPr"))
        if source_r_pr is not None:
            first_r_pr = deepcopy(source_r_pr)
    for child in list(p):
        if child is not p_pr:
            p.remove(child)
    run = OxmlElement("w:r")
    if first_r_pr is not None:
        run.append(first_r_pr)
    text = OxmlElement("w:t")
    if new_text.startswith(" ") or new_text.endswith(" "):
        text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    text.text = new_text
    run.append(text)
    p.append(run)


def replace_cell_text(cell, new_text: str) -> None:
    paragraph = cell.paragraphs[0]
    p = paragraph._p
    p_pr = p.find(qn("w:pPr"))
    first_r_pr = None
    first_run = p.find(qn("w:r"))
    if first_run is not None:
        source_r_pr = first_run.find(qn("w:rPr"))
        if source_r_pr is not None:
            first_r_pr = deepcopy(source_r_pr)
    for child in list(p):
        if child is not p_pr:
            p.remove(child)
    run = OxmlElement("w:r")
    if first_r_pr is not None:
        run.append(first_r_pr)
    text = OxmlElement("w:t")
    text.text = new_text
    run.append(text)
    p.append(run)
    for extra in list(cell.paragraphs)[1:]:
        cell._tc.remove(extra._p)


def format_value(value, number_format: str) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value
    if isinstance(value, bool):
        return "Yes" if value else "No"
    if isinstance(value, int):
        return f"{value:,}" if "#,##0" in number_format else str(value)
    if isinstance(value, float):
        if "%" in number_format:
            decimals = 1 if "0.0%" in number_format else 0
            rendered = f"{abs(value) * 100:.{decimals}f}%"
            if value > 0 and number_format.startswith("+"):
                return "+" + rendered
            if value < 0:
                return "−" + rendered
            return rendered
        if "0.000" in number_format:
            return f"{value:.3f}"
        if "0.00" in number_format:
            return f"{value:.2f}"
        if "0.0" in number_format:
            return f"{value:,.1f}" if "#,##" in number_format else f"{value:.1f}"
        if "#,##0" in number_format:
            return f"{value:,.0f}"
        return f"{value:g}"
    return str(value)


def workbook_matrix(path: Path, *, omit_title: bool) -> list[list[str]]:
    workbook = load_workbook(path, data_only=True)
    sheet = workbook.worksheets[0]
    start_row = 2 if omit_title else 1
    return [
        [format_value(cell.value, cell.number_format) for cell in row]
        for row in sheet.iter_rows(min_row=start_row, max_row=sheet.max_row, max_col=sheet.max_column)
    ]


def find_unique_table(document: Document, first_text: str, columns: int | None = None):
    matches = [
        table for table in document.tables
        if table.cell(0, 0).text == first_text and (columns is None or len(table.columns) == columns)
    ]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one table {first_text!r}/{columns}, found {len(matches)}")
    return matches[0]


def update_table_in_place(table, matrix: list[list[str]], label: str) -> None:
    if (len(table.rows), len(table.columns)) != (len(matrix), len(matrix[0])):
        raise RuntimeError(
            f"{label} shape mismatch: DOCX {(len(table.rows), len(table.columns))}, "
            f"workbook {(len(matrix), len(matrix[0]))}"
        )
    for row_index, values in enumerate(matrix):
        for column_index, value in enumerate(values):
            replace_cell_text(table.cell(row_index, column_index), value)


def build_from_workbook(document: Document, path: Path, widths: list[float]):
    matrix = workbook_matrix(path, omit_title=False)
    title = matrix[0][0]
    headers = matrix[1]
    rows = matrix[2:]
    return build_table(document, title, headers, rows, widths, set(range(1, len(headers))))


def build_b7(document: Document, network_path: Path, service_path: Path):
    network = pd.read_csv(network_path)
    service = pd.read_csv(service_path)
    labels = [("strict_joint", "Strict joint"), ("central", "Central"), ("permissive_joint", "Permissive joint")]
    rows: list[list[str]] = []
    service_names = ["Shelter", "Fire service", "Municipal facility", "Emergency water"]
    for key, label in labels:
        net = network.loc[(network["transfer_setting"] == key) & (network["scenario"] == "Heavy")]
        if len(net) != 1:
            raise RuntimeError(f"Expected one Heavy network row for {key}")
        row = net.iloc[0]
        svc = service.loc[service["transfer_setting"] == key].set_index("service_class")
        if any(name not in svc.index for name in service_names):
            raise RuntimeError(f"Missing service row for {key}")
        rows.append([
            label,
            f"{int(row['candidate_road_sections']):,}",
            f"{row['expected_isolated_population_mean']:,.1f}",
            f"{row['expected_isolated_population_age65_mean']:,.1f}",
            f"{svc.loc['Shelter', 'expected_service_loss_population_mean']:,.1f}",
            f"{svc.loc['Fire service', 'expected_service_loss_population_mean']:,.1f}",
            f"{svc.loc['Municipal facility', 'expected_service_loss_population_mean']:,.1f}",
            f"{svc.loc['Emergency water', 'expected_service_loss_population_mean']:,.1f}",
        ])
    headers = [
        "Transfer setting", "Heavy candidate roads", "Expected disconnected population",
        "Expected disconnected population age 65+", "Shelter loss", "Fire-service loss",
        "Municipal-facility loss", "Emergency-water sensitivity",
    ]
    return build_table(
        document,
        "Table B7. Downstream slope-to-road transfer bounds",
        headers,
        rows,
        [1.10, 0.95, 1.05, 1.18, 0.85, 0.90, 1.02, 1.15],
        {1, 2, 3, 4, 5, 6, 7},
    )


def apply_update(args, dry_run: bool) -> dict[str, object]:
    docx = args.docx
    before_hash = sha256(docx)
    before_members = package_members(docx)
    document = Document(docx)

    protected_titles = [
        "Table B6a. Slope-to-road parameter specifications",
        "Table B6b. Road-ranking and episode-correspondence sensitivity",
        "Appendix Table B9. Service-destination estimand and rerouting comparison",
    ]
    protected_before = {title: table_hash(find_unique_table(document, title)) for title in protected_titles}

    for old, new in PARAGRAPH_REPLACEMENTS.items():
        matches = [p for p in document.paragraphs if paragraph_visible_text(p) == old]
        if len(matches) != 1:
            raise RuntimeError(f"Expected one paragraph match, found {len(matches)} for {old[:70]!r}")
        replace_paragraph_text(matches[0], new)

    a1 = find_unique_table(document, "Analytical Data Layer", 10)
    role_matches = [cell for row in a1.rows for cell in row.cells if cell.text == OLD_A1_ROLE]
    if len(role_matches) != 1:
        raise RuntimeError(f"Expected one A1 role cell, found {len(role_matches)}")
    replace_cell_text(role_matches[0], NEW_A1_ROLE)

    in_place_specs = [
        ("Sensitivity Setting", 10, args.b3_workbook, "B3"),
        ("Comparator", 8, args.b4_workbook, "B4"),
        ("Admin Area Code", 16, args.c1_workbook, "C1"),
        ("Priority Rank", 12, args.c2_workbook, "C2"),
        ("Priority Rank", 11, args.c3_workbook, "C3"),
    ]
    for first_text, columns, workbook, label in in_place_specs:
        update_table_in_place(
            find_unique_table(document, first_text, columns),
            workbook_matrix(workbook, omit_title=True),
            label,
        )

    old_replacement_tables = {
        title: find_unique_table(document, title)
        for title in [
            "Table B5. Baseline-Threshold Comparison and Rainfall-Parameter Sensitivity",
            "Table B7. Downstream slope-to-road transfer bounds",
            "Table B8. Spatial closure-dependence sensitivity",
        ]
    }
    replacement_specs = [
        (
            "Table B5. Baseline-Threshold Comparison and Rainfall-Parameter Sensitivity",
            build_from_workbook(document, args.b5_workbook, [1.08, 1.48, 1.10, 1.30, 1.02, 1.10, 1.52]),
        ),
        (
            "Table B7. Downstream slope-to-road transfer bounds",
            build_b7(document, args.b7_network_csv, args.b7_service_csv),
        ),
        (
            "Table B8. Spatial closure-dependence sensitivity",
            build_from_workbook(
                document,
                args.b8_workbook,
                [0.67, 0.82, 0.58, 0.43, 0.92, 0.92, 0.72, 0.68, 0.72, 0.92, 0.72, 0.92],
            ),
        ),
    ]
    for old_title, new_table in replacement_specs:
        old_table = old_replacement_tables[old_title]
        old_table._tbl.getparent().replace(old_table._tbl, new_table._tbl)

    if any(table.cell(0, 0).text.startswith("Table B10.") for table in document.tables):
        raise RuntimeError("Table B10 already exists")
    b10 = build_from_workbook(
        document,
        args.b10_workbook,
        [1.45, 0.72, 0.62, 0.74, 0.70, 0.78, 1.35, 1.42, 1.02, 1.02],
    )
    b9 = find_unique_table(document, protected_titles[2])
    page_break = page_break_paragraph()
    spacer = OxmlElement("w:p")
    b9._tbl.addnext(page_break)
    page_break.addnext(b10._tbl)
    b10._tbl.addnext(spacer)

    with tempfile.NamedTemporaryFile(prefix=f".{docx.name}.", suffix=".tmp.docx", dir=docx.parent, delete=False) as handle:
        staged = Path(handle.name)
    try:
        document.save(staged)
        verified = Document(staged)
        for new in PARAGRAPH_REPLACEMENTS.values():
            if sum(paragraph_visible_text(p) == new for p in verified.paragraphs) != 1:
                raise RuntimeError(f"Paragraph verification failed: {new[:70]!r}")
        if sum(cell.text == NEW_A1_ROLE for row in find_unique_table(verified, "Analytical Data Layer", 10).rows for cell in row.cells) != 1:
            raise RuntimeError("A1 role verification failed")

        expected_shapes = {
            "Sensitivity Setting": (22, 10),
            "Comparator": (85, 8),
            "Table B5. Baseline-Threshold Comparison and Rainfall-Parameter Sensitivity": (20, 7),
            "Table B7. Downstream slope-to-road transfer bounds": (5, 8),
            "Table B8. Spatial Closure-Dependence Sensitivity": (20, 12),
            "Table B10. Emergency-road-backbone target sensitivity": (11, 10),
            "Admin Area Code": (50, 16),
        }
        for title, shape in expected_shapes.items():
            table = find_unique_table(verified, title)
            actual = (len(table.rows), len(table.columns))
            if actual != shape:
                raise RuntimeError(f"{title} shape {actual}, expected {shape}")
        if (len(find_unique_table(verified, "Priority Rank", 12).rows), 12) != (31, 12):
            raise RuntimeError("C2 shape verification failed")
        if (len(find_unique_table(verified, "Priority Rank", 11).rows), 11) != (31, 11):
            raise RuntimeError("C3 shape verification failed")
        b10_check = find_unique_table(verified, "Table B10. Emergency-road-backbone target sensitivity", 10)
        if [b10_check.cell(row, 8).text for row in (4, 7, 10)] != ["Not comparable"] * 3:
            raise RuntimeError("B10 audit-comparator labeling verification failed")
        protected_after = {title: table_hash(find_unique_table(verified, title)) for title in protected_titles}
        if protected_after != protected_before:
            raise RuntimeError("A protected B6/B9 table changed")

        after_members = package_members(staged)
        changed_members = sorted(
            name for name in set(before_members) | set(after_members)
            if before_members.get(name) != after_members.get(name)
        )
        forbidden = [name for name in changed_members if name.startswith("word/") and name != "word/document.xml"]
        if forbidden:
            raise RuntimeError(f"Protected Word package members changed: {forbidden}")
        receipt: dict[str, object] = {
            "status": "dry-run" if dry_run else "applied",
            "sha256_before": before_hash,
            "sha256_after": sha256(staged),
            "changed_package_members": changed_members,
            "paragraphs_replaced": len(PARAGRAPH_REPLACEMENTS),
            "a1_role_updated": True,
            "tables_updated": ["B3", "B4", "B5", "B7", "B8", "C1", "C2", "C3"],
            "table_b10_added": True,
            "protected_table_hashes_unchanged": protected_after == protected_before,
            "backup": None,
        }
        if dry_run:
            return receipt

        backup_dir = docx.parent / ".kila-backups"
        backup_dir.mkdir(parents=True, exist_ok=True)
        stamp = dt.datetime.now().strftime("%Y%m%dT%H%M%S%f")
        backup = backup_dir / f"Appendix.before-r3c6.{stamp}.docx"
        shutil.copy2(docx, backup)
        if sha256(backup) != before_hash:
            raise RuntimeError("Backup hash mismatch")
        os.replace(staged, docx)
        receipt["backup"] = str(backup)
        if sha256(docx) != receipt["sha256_after"]:
            raise RuntimeError("Final Appendix hash differs from the validated staged file")
        return receipt
    finally:
        if staged.exists():
            staged.unlink()


def parse_args():
    parser = argparse.ArgumentParser()
    parser.add_argument("--docx", type=Path, default=Path("Rev/revision/Appendix.docx"))
    parser.add_argument("--b3-workbook", type=Path, default=Path("data/results/tables/Table_intervention_portfolios.xlsx"))
    parser.add_argument("--b4-workbook", type=Path, default=Path("data/results/tables/Table_comparator_robustness.xlsx"))
    parser.add_argument("--b5-workbook", type=Path, required=True, help="Recalculated data-only-readable B5 workbook")
    parser.add_argument("--b7-network-csv", type=Path, default=Path("data/exp/revision/reviewer-2-comment-5/downstream_network_sensitivity.csv"))
    parser.add_argument("--b7-service-csv", type=Path, default=Path("data/exp/revision/reviewer-2-comment-5/downstream_service_sensitivity.csv"))
    parser.add_argument("--b8-workbook", type=Path, default=Path("data/results/tables/Table_spatial_closure_dependence_sensitivity.xlsx"))
    parser.add_argument("--b10-workbook", type=Path, default=Path("data/results/tables/Table_emergency_backbone_target_sensitivity.xlsx"))
    parser.add_argument("--c1-workbook", type=Path, default=Path("data/results/tables/Table_municipality_isolation_and_service_loss_summary.xlsx"))
    parser.add_argument("--c2-workbook", type=Path, default=Path("data/results/tables/Table_priority_road_sections.xlsx"))
    parser.add_argument("--c3-workbook", type=Path, default=Path("data/results/tables/Table_high_isolation_risk_communities.xlsx"))
    parser.add_argument("--apply", action="store_true")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    print(json.dumps(apply_update(args, dry_run=not args.apply), indent=2, ensure_ascii=False))


if __name__ == "__main__":
    main()
