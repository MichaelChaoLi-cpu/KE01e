#!/usr/bin/env python3
"""Apply the approved Reviewer 1 Comment 2 Appendix text updates."""

from __future__ import annotations

import argparse
from copy import deepcopy
import datetime as dt
import hashlib
import json
import os
from pathlib import Path
import shutil
import tempfile
import zipfile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


REPLACEMENTS = {
    (
        "Simulation-size and network-target checks bound the disconnection results. "
        "Changing the number of draws from 500 to 2,000 with one common seed produces "
        "a 0.006 difference in the 95th-percentile community isolation frequency. The "
        "Primary Emergency Road backbone contains 2,562 target roots and yields 1,063.6 "
        "expected disconnected residents under Heavy rainfall. Expanding the target to "
        "2,977 Primary-plus-Secondary Emergency Road roots yields 992.7 residents, with "
        "community-frequency Spearman correlation 0.964 and top-30 burden overlap 90.0% "
        "relative to the primary definition. The former coast-inclusive boundary proxy "
        "is reported only as an audit comparator and reproduces the prior 1,121.7-resident "
        "result. Municipality-wide Yatsushiro assignments of 0.70 and 0.80 bound the revised "
        "primary result between 1,016.6 and 1,118.7, while alternative closure mappings "
        "produce the wider rounded range of 343–2,057 residents. Closure mapping is therefore "
        "a larger source of magnitude uncertainty than Monte Carlo convergence or the tested "
        "emergency-road target definition."
    ): (
        "Simulation-size and network-target checks bound the disconnection results. "
        "Changing the number of draws from 500 to 2,000 with one common seed produces "
        "a 0.006 difference in the 95th-percentile community isolation frequency. The "
        "Primary Emergency Road backbone contains 2,562 target roots and yields 1,063.6 "
        "expected disconnected residents under Heavy rainfall. Expanding the target to "
        "2,977 Primary-plus-Secondary Emergency Road roots yields 992.7 residents, with "
        "community-frequency Spearman correlation 0.964 and top-30 burden overlap 90.0% "
        "relative to the primary definition. The former coast-inclusive boundary proxy "
        "is reported only as an audit comparator and reproduces the prior 1,121.7-resident "
        "result. Municipality-wide Yatsushiro assignments of 0.70 and 0.80 bound the revised "
        "primary result between 1,016.6 and 1,118.7 residents (−4.4% to +5.2% relative to the "
        "0.75 midpoint); community-frequency rank correlations remain 0.989–1.000 and Top-30 "
        "population-burden overlap is 93.3%. Alternative closure mappings produce the wider "
        "rounded range of 343–2,057 residents. Closure mapping is therefore a larger source "
        "of magnitude uncertainty than Monte Carlo convergence or the tested emergency-road "
        "target definition."
    ),
    (
        "Service sensitivity retains only baseline-reachable community-destination pairs and "
        "does not convert pre-existing disconnection into disruption-induced loss. "
        "Municipality-wide Yatsushiro assignments of 0.70 and 0.80 bound Heavy-scenario "
        "affected populations at 563.1–647.2 for shelters, 7,636.2–8,466.1 for emergency water, "
        "1,178.2–1,353.8 for fire services, and 885.6–969.4 for municipal facilities. "
        "Emergency-water routing is restricted to 10 resolved point locations among 36 "
        "announcements, so its affected-population and excess-time estimates apply only to "
        "that destination set. Additional valid destinations could reduce simulated loss, "
        "while differences between announced and operational points could change it in either "
        "direction."
    ): (
        "Service sensitivity retains only baseline-reachable community-destination pairs and "
        "does not convert pre-existing disconnection into disruption-induced loss. "
        "Municipality-wide Yatsushiro assignments of 0.70 and 0.80 bound Heavy-scenario "
        "affected populations at 570.5–657.0 for shelters, 7,582.5–8,238.8 for emergency water, "
        "1,215.1–1,389.7 for fire services, and 898.6–984.2 for municipal facilities. Across "
        "the three primary service classes, changes relative to the 0.75 midpoint remain within "
        "−6.5% to +8.0%, frequency rank correlations are 0.990–1.000, and Top-30 burden overlap "
        "is 86.7%–93.3%. Emergency-water routing is restricted to 10 resolved point locations "
        "among 36 announcements, so its affected-population and excess-time estimates apply "
        "only to that destination set. Additional valid destinations could reduce simulated "
        "loss, while differences between announced and operational points could change it in "
        "either direction."
    ),
}


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def package_members(path: Path) -> dict[str, bytes]:
    with zipfile.ZipFile(path) as archive:
        bad = archive.testzip()
        if bad:
            raise RuntimeError(f"Corrupt DOCX member: {bad}")
        return {name: archive.read(name) for name in archive.namelist()}


def visible_text(paragraph) -> str:
    tags = {qn("w:t"), qn("w:delText"), qn("m:t")}
    return "".join(element.text or "" for element in paragraph._p.iter() if element.tag in tags)


def table_hashes(document: Document) -> list[str]:
    return [hashlib.sha256(table._tbl.xml.encode("utf-8")).hexdigest() for table in document.tables]


def bookmark_fingerprint(paragraph) -> str:
    parts: list[str] = []
    for element in paragraph._p:
        if element.tag in {qn("w:bookmarkStart"), qn("w:bookmarkEnd")}:
            attrs = sorted((key, value) for key, value in element.attrib.items())
            parts.append(f"{element.tag}:{attrs}")
    return hashlib.sha256("\n".join(parts).encode("utf-8")).hexdigest()


def replace_visible_content(paragraph, new_text: str) -> None:
    p = paragraph._p
    p_pr = p.find(qn("w:pPr"))
    first_run = p.find(qn("w:r"))
    r_pr = None
    if first_run is not None and first_run.find(qn("w:rPr")) is not None:
        r_pr = deepcopy(first_run.find(qn("w:rPr")))

    protected = {qn("w:pPr"), qn("w:bookmarkStart"), qn("w:bookmarkEnd")}
    for child in list(p):
        if child.tag not in protected:
            p.remove(child)

    run = OxmlElement("w:r")
    if r_pr is not None:
        run.append(r_pr)
    text = OxmlElement("w:t")
    text.text = new_text
    run.append(text)

    bookmark_ends = [index for index, child in enumerate(p) if child.tag == qn("w:bookmarkEnd")]
    insertion_index = min(bookmark_ends) if bookmark_ends else len(p)
    p.insert(insertion_index, run)
    if p_pr is not None and p[0] is not p_pr:
        raise RuntimeError("Paragraph properties moved during replacement")


def update(docx: Path, dry_run: bool) -> dict[str, object]:
    before_hash = sha256(docx)
    before_members = package_members(docx)
    document = Document(docx)
    before_tables = table_hashes(document)
    before_bookmarks: dict[str, str] = {}

    for old, new in REPLACEMENTS.items():
        matches = [paragraph for paragraph in document.paragraphs if visible_text(paragraph) == old]
        if len(matches) != 1:
            raise RuntimeError(f"Expected one exact paragraph match, found {len(matches)} for {old[:80]!r}")
        paragraph = matches[0]
        before_bookmarks[old] = bookmark_fingerprint(paragraph)
        replace_visible_content(paragraph, new)
        if visible_text(paragraph) != new:
            raise RuntimeError("In-memory replacement verification failed")
        if bookmark_fingerprint(paragraph) != before_bookmarks[old]:
            raise RuntimeError("Bookmark fingerprint changed during replacement")

    with tempfile.NamedTemporaryFile(
        prefix=f".{docx.name}.", suffix=".tmp.docx", dir=docx.parent, delete=False
    ) as handle:
        staged = Path(handle.name)
    try:
        document.save(staged)
        verified = Document(staged)
        for new in REPLACEMENTS.values():
            if sum(visible_text(paragraph) == new for paragraph in verified.paragraphs) != 1:
                raise RuntimeError("Saved replacement paragraph verification failed")
        if table_hashes(verified) != before_tables:
            raise RuntimeError("An Appendix table changed during paragraph-only update")
        after_members = package_members(staged)
        changed = sorted(
            name for name in set(before_members) | set(after_members)
            if before_members.get(name) != after_members.get(name)
        )
        forbidden = [
            name for name in changed
            if name.startswith("word/") and name != "word/document.xml"
        ]
        if forbidden:
            raise RuntimeError(f"Protected Word members changed: {forbidden}")
        receipt: dict[str, object] = {
            "status": "dry-run" if dry_run else "applied",
            "sha256_before": before_hash,
            "sha256_after": sha256(staged),
            "changed_package_members": changed,
            "paragraphs_replaced": len(REPLACEMENTS),
            "tables_unchanged": True,
            "bookmarks_preserved": True,
            "backup": None,
        }
        if dry_run:
            return receipt
        backup_dir = docx.parent / ".kila-backups"
        backup_dir.mkdir(parents=True, exist_ok=True)
        stamp = dt.datetime.now().strftime("%Y%m%dT%H%M%S%f")
        backup = backup_dir / f"Appendix.before-r1c2.{stamp}.docx"
        shutil.copy2(docx, backup)
        if sha256(backup) != before_hash:
            raise RuntimeError("Backup hash mismatch")
        os.replace(staged, docx)
        receipt["backup"] = str(backup)
        return receipt
    finally:
        if staged.exists():
            staged.unlink()


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--docx", type=Path, default=Path("Rev/revision/Appendix.docx"))
    parser.add_argument("--apply", action="store_true")
    args = parser.parse_args()
    print(json.dumps(update(args.docx, dry_run=not args.apply), indent=2))


if __name__ == "__main__":
    main()
