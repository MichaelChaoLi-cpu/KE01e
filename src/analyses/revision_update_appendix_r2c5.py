#!/usr/bin/env python3
"""Apply the approved Reviewer 2 Comment 5 Appendix B update transactionally."""

from __future__ import annotations

import argparse
import datetime as dt
import hashlib
import json
import os
from pathlib import Path
import shutil
import tempfile
import zipfile

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ANCHOR = (
    "The rainfall-scenario scores have near-unity rank agreement, so they support "
    "relative screening and magnitude contrasts but not scenario-specific spatial "
    "reprioritization or an identified population failure rate."
)
HEADING = "Slope-to-Road Transfer Specification and Sensitivity"
METHOD_TEXT = (
    "The road-transfer grid contains 752 × 950 cells with approximate midpoint "
    "dimensions of 144 × 170 m. Each line component is sampled at fractions 0.20, "
    "0.50, and 0.80. The central influence set searches all noncentral offsets "
    "within ±3 cells, requires at least 10 m positive relief and alignment cosine "
    "of 0.20, and weights candidates by exponential distance decay with a 2.5-cell "
    "e-folding length, clipped alignment, and relief scaled by 100 m with limits "
    "0.20 and 1.00. Table B6 lists the central setting and all prespecified "
    "alternatives. These values define a reproducible regional-screening rule and "
    "are not calibrated runout parameters."
)
RESULTS_TEXT = (
    "Across 15 specifications, the strict joint boundary gives the lowest "
    "supported-road rank correlation (0.676) and top-1% overlap (0.428), while "
    "every matched-concordance estimate remains above 0.50. The main instability "
    "arises from neighborhood reach and relief-based support rather than from the "
    "continuous distance, alignment, or relief weights alone. Table B7 propagates "
    "the joint boundaries using each setting's own Heavy 85th-percentile candidate "
    "set and 99.5th-percentile closure-mapping upper bound. Heavy expected isolated "
    "population spans 523.5–2,256.3 residents around the central 1,121.7, so exact "
    "priorities and consequence magnitude remain conditional on the transfer "
    "specification."
)
OLD_SENTENCE = (
    "Thus, location screening is comparatively robust, whereas consequence "
    "magnitude remains sensitive to the rainfall parameterization."
)
NEW_SENTENCE = (
    "Thus, location screening is comparatively robust to the tested "
    "rainfall-window and γ choices, whereas consequence magnitude remains sensitive "
    "to the rainfall parameterization; Appendix Tables B6 and B7 separately show "
    "that the slope-to-road influence-set boundaries materially affect exact road "
    "priorities and downstream magnitude."
)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def package_members(path: Path) -> dict[str, bytes]:
    with zipfile.ZipFile(path) as archive:
        bad = archive.testzip()
        if bad:
            raise RuntimeError(f"Corrupt DOCX member: {bad}")
        return {name: archive.read(name) for name in archive.namelist()}


def paragraph_with_exact_text(document: Document, text: str):
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text == text]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one exact paragraph match, found {len(matches)}")
    return matches[0]


def paragraph_containing_text(document: Document, text: str):
    matches = [paragraph for paragraph in document.paragraphs if text in paragraph.text]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph containing anchor, found {len(matches)}")
    return matches[0]


def insert_paragraph_after(previous, text: str, style: str):
    paragraph = Document().add_paragraph(text, style=style)
    element = paragraph._p
    previous.addnext(element)
    return element


def set_cell_text(
    cell,
    text: str,
    *,
    bold: bool,
    size: float,
    align,
    color: str = "000000",
    shading: str | None = None,
    no_wrap: bool = False,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.right_indent = Pt(0)
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "left", "bottom", "right"):
        margin = tc_mar.find(qn(f"w:{side}"))
        if margin is None:
            margin = OxmlElement(f"w:{side}")
            tc_mar.append(margin)
        margin.set(qn("w:w"), "35")
        margin.set(qn("w:type"), "dxa")
    if no_wrap:
        marker = tc_pr.find(qn("w:noWrap"))
        if marker is None:
            tc_pr.append(OxmlElement("w:noWrap"))
    if shading:
        shd = tc_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tc_pr.append(shd)
        shd.set(qn("w:fill"), shading)
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "B7C9D6")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:cantSplit")
    tr_pr.append(marker)


def page_break_paragraph() -> OxmlElement:
    paragraph = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    page_break = OxmlElement("w:pageBreakBefore")
    p_pr.append(page_break)
    paragraph.append(p_pr)
    return paragraph


def build_table(
    document: Document,
    title: str,
    headers: list[str],
    rows: list[list[str]],
    widths: list[float],
    nowrap_columns: set[int],
):
    table = document.add_table(rows=2 + len(rows), cols=len(headers))
    table.style = "Normal Table"
    table.autofit = False
    table_pr = table._tbl.tblPr
    layout = table_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    for column, width in zip(table.columns, widths):
        column.width = Inches(width)
        for cell in column.cells:
            cell.width = Inches(width)
    title_cell = table.cell(0, 0).merge(table.cell(0, len(headers) - 1))
    set_cell_text(
        title_cell,
        title,
        bold=True,
        size=10.0,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        color="1F4E78",
        shading="D9EAF7",
    )
    for column, header in enumerate(headers):
        cell = table.cell(1, column)
        cell.width = Inches(widths[column])
        set_cell_text(
            cell,
            header,
            bold=True,
            size=6.5,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            color="FFFFFF",
            shading="1F4E78",
        )
    set_repeat_table_header(table.rows[0])
    set_repeat_table_header(table.rows[1])
    for row_index, values in enumerate(rows, start=2):
        for column, value in enumerate(values):
            cell = table.cell(row_index, column)
            cell.width = Inches(widths[column])
            set_cell_text(
                cell,
                value,
                bold=False,
            size=6.3,
                align=(
                    WD_ALIGN_PARAGRAPH.LEFT
                    if column == 0
                    else WD_ALIGN_PARAGRAPH.CENTER
                ),
                shading="F4F7F9" if row_index % 2 == 0 else "FFFFFF",
                no_wrap=column in nowrap_columns,
            )
        set_row_cant_split(table.rows[row_index])
    return table


def table_b6_rows(data_dir: Path) -> list[list[str]]:
    parameters = pd.read_csv(data_dir / "parameter_specifications.csv")
    road = pd.read_csv(data_dir / "road_score_sensitivity.csv")
    validation = pd.read_csv(data_dir / "matched_validation_sensitivity.csv")
    frame = (
        parameters.merge(road, left_on="key", right_on="specification")
        .merge(validation, on="specification")
        .set_index("key")
    )
    order = [
        ("central", "Central"),
        ("radius_2", "Radius 2 cells"),
        ("radius_4", "Radius 4 cells"),
        ("relief_5", "Relief 5 m"),
        ("relief_20", "Relief 20 m"),
        ("alignment_0", "Alignment 0.00"),
        ("alignment_05", "Alignment 0.50"),
        ("decay_15", "Distance 1.5 cells"),
        ("decay_40", "Distance 4.0 cells"),
        ("relief_scale_50", "Relief scale 50 m"),
        ("relief_scale_150", "Relief scale 150 m"),
        ("midpoint", "Midpoint only"),
        ("five_points", "Five points"),
        ("strict_joint", "Strict joint"),
        ("permissive_joint", "Permissive joint"),
    ]
    rows: list[list[str]] = []
    for key, label in order:
        row = frame.loc[key]
        rows.append(
            [
                label,
                f"{int(row.radius_cells)}",
                f"{row.minimum_relief_m:.0f}",
                f"{row.minimum_alignment_cosine:.2f}",
                f"{row.distance_efold_cells:.1f}",
                f"{row.relief_scale_m:.0f}",
                str(row.sample_fractions).replace(",", ", "),
                f"{int(row.supported_road_sections):,} ({100 * row.support_fraction:.1f}%)",
                f"{row.spearman_union_supported:.3f}",
                f"{100 * row.top1_overlap_with_central:.1f}%",
                f"{100 * row.official_heavy_candidate_overlap_with_central:.1f}%",
                f"{row.matched_concordance:.3f} ({row.bootstrap_ci_low:.3f}–{row.bootstrap_ci_high:.3f})",
            ]
        )
    return rows


def table_b7_rows(data_dir: Path) -> list[list[str]]:
    network = pd.read_csv(data_dir / "downstream_network_sensitivity.csv")
    network = network[network.scenario == "Heavy"].set_index("transfer_setting")
    services = pd.read_csv(data_dir / "downstream_service_sensitivity.csv")
    services = services.pivot(
        index="transfer_setting",
        columns="service_class",
        values="expected_service_loss_population_mean",
    )
    labels = [
        ("strict_joint", "Strict joint"),
        ("central", "Central"),
        ("permissive_joint", "Permissive joint"),
    ]
    rows = []
    for key, label in labels:
        n = network.loc[key]
        s = services.loc[key]
        rows.append(
            [
                label,
                f"{int(n.candidate_road_sections):,}",
                f"{n.expected_isolated_population_mean:,.1f}",
                f"{n.expected_isolated_population_age65_mean:,.1f}",
                f"{s['Shelter']:,.1f}",
                f"{s['Fire service']:,.1f}",
                f"{s['Municipal facility']:,.1f}",
                f"{s['Emergency water']:,.1f}",
            ]
        )
    return rows


def apply_update(docx: Path, data_dir: Path, dry_run: bool) -> dict[str, object]:
    before_hash = sha256(docx)
    protected_before = package_members(docx)
    document = Document(docx)
    if any(paragraph.text == HEADING for paragraph in document.paragraphs):
        raise RuntimeError("Appendix subsection already exists")
    if any(table.cell(0, 0).text.startswith("Table B6") for table in document.tables):
        raise RuntimeError("Appendix Table B6 already exists")

    anchor = paragraph_containing_text(document, ANCHOR)
    cursor = anchor._p
    cursor = insert_paragraph_after(cursor, HEADING, "Heading 2")
    cursor = insert_paragraph_after(cursor, METHOD_TEXT, "Normal")
    insert_paragraph_after(cursor, RESULTS_TEXT, "Normal")

    paragraph9 = [paragraph for paragraph in document.paragraphs if OLD_SENTENCE in paragraph.text]
    if len(paragraph9) != 1 or len(paragraph9[0].runs) != 1:
        raise RuntimeError("Part 9 target is not one exact single-run paragraph")
    paragraph9[0].runs[0].text = paragraph9[0].runs[0].text.replace(
        OLD_SENTENCE, NEW_SENTENCE
    )

    b6_full_rows = table_b6_rows(data_dir)
    b6a_headers = [
        "Specification",
        "Radius\n(cells)",
        "Min. relief\n(m)",
        "Min. align.\ncosine",
        "Distance e-fold\n(cells)",
        "Relief scale\n(m)",
        "Road samples\n(fractions)",
    ]
    b6a = build_table(
        document,
        "Table B6a. Slope-to-road parameter specifications",
        b6a_headers,
        [row[:7] for row in b6_full_rows],
        [1.50, 0.70, 0.85, 1.00, 1.10, 0.85, 1.60],
        {1, 2, 3, 4, 5},
    )
    b6b_headers = [
        "Specification",
        "Supported\nroads",
        "Spearman\nvs central",
        "Top-1%\noverlap",
        "Heavy set\noverlap",
        "Matched concord.\n(95% CI)",
    ]
    b6b = build_table(
        document,
        "Table B6b. Road-ranking and matched-evidence sensitivity",
        b6b_headers,
        [[row[0], *row[7:]] for row in b6_full_rows],
        [1.65, 1.35, 1.20, 1.05, 1.15, 2.10],
        {2, 3, 4},
    )
    b7_headers = [
        "Transfer setting",
        "Heavy candidate\nroads",
        "Expected isolated\npopulation",
        "Expected isolated\npopulation age 65+",
        "Shelter loss",
        "Fire-service loss",
        "Municipal-facility\nloss",
        "Emergency-water\nsensitivity",
    ]
    b7 = build_table(
        document,
        "Table B7. Downstream slope-to-road transfer bounds",
        b7_headers,
        table_b7_rows(data_dir),
        [1.10, 0.95, 1.05, 1.18, 0.85, 0.90, 1.02, 1.15],
        {1, 2, 3, 4, 5, 6, 7},
    )
    c1_title = paragraph_with_exact_text(
        document, "Appendix Table C1. Municipality isolation and service-loss summary"
    )._p
    c1_title.addprevious(page_break_paragraph())
    c1_title.addprevious(b6a._tbl)
    c1_title.addprevious(page_break_paragraph())
    c1_title.addprevious(b6b._tbl)
    c1_title.addprevious(page_break_paragraph())
    c1_title.addprevious(b7._tbl)
    c1_title.addprevious(OxmlElement("w:p"))

    with tempfile.NamedTemporaryFile(
        prefix=f".{docx.name}.", suffix=".tmp.docx", dir=docx.parent, delete=False
    ) as handle:
        temp_path = Path(handle.name)
    try:
        document.save(temp_path)
        check = Document(temp_path)
        texts = [paragraph.text for paragraph in check.paragraphs]
        if texts.count(HEADING) != 1 or OLD_SENTENCE in "\n".join(texts):
            raise RuntimeError("Appendix text verification failed")
        if NEW_SENTENCE not in "\n".join(texts):
            raise RuntimeError("Appendix Part 9 replacement verification failed")
        tables = {table.cell(0, 0).text: table for table in check.tables}
        b6a_check = tables.get("Table B6a. Slope-to-road parameter specifications")
        b6b_check = tables.get("Table B6b. Road-ranking and matched-evidence sensitivity")
        b7_check = tables.get("Table B7. Downstream slope-to-road transfer bounds")
        if b6a_check is None or len(b6a_check.rows) != 17 or len(b6a_check.columns) != 7:
            raise RuntimeError("Appendix Table B6a shape verification failed")
        if b6b_check is None or len(b6b_check.rows) != 17 or len(b6b_check.columns) != 6:
            raise RuntimeError("Appendix Table B6b shape verification failed")
        if b7_check is None or len(b7_check.rows) != 5 or len(b7_check.columns) != 8:
            raise RuntimeError("Appendix Table B7 shape verification failed")
        protected_after = package_members(temp_path)
        changed_members = sorted(
            name
            for name in set(protected_before) | set(protected_after)
            if protected_before.get(name) != protected_after.get(name)
        )
        forbidden = [
            name
            for name in changed_members
            if name.startswith("word/") and name != "word/document.xml"
        ]
        if forbidden:
            raise RuntimeError(f"Protected Word package members changed: {forbidden}")
        receipt: dict[str, object] = {
            "status": "dry-run" if dry_run else "applied",
            "sha256_before": before_hash,
            "sha256_after": sha256(temp_path),
            "changed_package_members": changed_members,
            "table_b6a_shape": [len(b6a_check.rows), len(b6a_check.columns)],
            "table_b6b_shape": [len(b6b_check.rows), len(b6b_check.columns)],
            "table_b7_shape": [len(b7_check.rows), len(b7_check.columns)],
            "backup": None,
        }
        if dry_run:
            return receipt
        backup_dir = docx.parent / ".kila-backups"
        backup_dir.mkdir(parents=True, exist_ok=True)
        stamp = dt.datetime.now().strftime("%Y%m%dT%H%M%S%f")
        backup = backup_dir / f"{docx.stem}.{stamp}.reviewer-2-comment-5.parts-08-09.docx"
        shutil.copy2(docx, backup)
        if sha256(backup) != before_hash:
            raise RuntimeError("Appendix backup hash mismatch")
        os.replace(temp_path, docx)
        receipt["backup"] = str(backup)
        if sha256(docx) != receipt["sha256_after"]:
            raise RuntimeError("Final Appendix hash differs from validated temporary file")
        return receipt
    finally:
        if temp_path.exists():
            temp_path.unlink()


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--docx", type=Path, required=True)
    parser.add_argument("--data-dir", type=Path, required=True)
    parser.add_argument("--dry-run", action="store_true")
    args = parser.parse_args()
    print(json.dumps(apply_update(args.docx, args.data_dir, args.dry_run), indent=2))


if __name__ == "__main__":
    main()
