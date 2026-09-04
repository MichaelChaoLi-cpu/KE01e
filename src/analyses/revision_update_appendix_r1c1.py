#!/usr/bin/env python3
"""Insert the approved Reviewer 1 Comment 1 Appendix support paragraph."""

from __future__ import annotations

import argparse
from copy import deepcopy
import datetime as dt
import hashlib
import json
import os
from pathlib import Path
import shutil
import tempfile
import zipfile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ANCHOR = (
    "Unresolved records remain in the audit but do not enter point-based routing "
    "or fine-location validation."
)

INSERTED = (
    "The emergency-water subset is geographically and institutionally concentrated. "
    "The 10 resolved records are Yatsushiro school-named locations matched exactly to "
    "the retained 2012 facility register; the unresolved records comprise 18 of 28 "
    "Yatsushiro announcements, all six Uki announcements, and both Hikawa announcements. "
    "By name class, resolution is 10/13 for schools and 0/14, 0/8, and 0/1 for community "
    "centres, government/disaster centres, and other locations, respectively. These counts "
    "establish non-uniform data support but cannot identify urban–rural missingness without "
    "coordinates."
)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def package_members(path: Path) -> dict[str, bytes]:
    with zipfile.ZipFile(path) as archive:
        bad = archive.testzip()
        if bad:
            raise RuntimeError(f"Corrupt DOCX member: {bad}")
        return {name: archive.read(name) for name in archive.namelist()}


def visible_text(paragraph) -> str:
    tags = {qn("w:t"), qn("w:delText"), qn("m:t")}
    return "".join(
        element.text or "" for element in paragraph._p.iter() if element.tag in tags
    )


def table_hashes(document: Document) -> list[str]:
    return [
        hashlib.sha256(table._tbl.xml.encode("utf-8")).hexdigest()
        for table in document.tables
    ]


def bookmark_fingerprint(document: Document) -> str:
    parts: list[str] = []
    for element in document.element.body.iter():
        if element.tag in {qn("w:bookmarkStart"), qn("w:bookmarkEnd")}:
            parts.append(f"{element.tag}:{sorted(element.attrib.items())}")
    return hashlib.sha256("\n".join(parts).encode("utf-8")).hexdigest()


def insert_after(anchor, text: str) -> None:
    new_paragraph = OxmlElement("w:p")
    p_pr = anchor._p.find(qn("w:pPr"))
    if p_pr is not None:
        new_paragraph.append(deepcopy(p_pr))

    source_run = anchor._p.find(qn("w:r"))
    if source_run is None:
        raise RuntimeError("Anchor paragraph has no ordinary run for style reuse.")
    run = OxmlElement("w:r")
    r_pr = source_run.find(qn("w:rPr"))
    if r_pr is not None:
        run.append(deepcopy(r_pr))
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    new_paragraph.append(run)
    anchor._p.addnext(new_paragraph)


def update(docx: Path, dry_run: bool) -> dict[str, object]:
    before_hash = sha256(docx)
    before_members = package_members(docx)
    document = Document(docx)
    before_tables = table_hashes(document)
    before_bookmarks = bookmark_fingerprint(document)

    if sum(visible_text(paragraph) == INSERTED for paragraph in document.paragraphs):
        raise RuntimeError("Approved paragraph is already present.")
    matches = [
        paragraph for paragraph in document.paragraphs if ANCHOR in visible_text(paragraph)
    ]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one anchor paragraph, found {len(matches)}.")
    insert_after(matches[0], INSERTED)

    with tempfile.NamedTemporaryFile(
        prefix=f".{docx.name}.", suffix=".tmp.docx", dir=docx.parent, delete=False
    ) as handle:
        staged = Path(handle.name)
    try:
        document.save(staged)
        verified = Document(staged)
        if sum(visible_text(paragraph) == INSERTED for paragraph in verified.paragraphs) != 1:
            raise RuntimeError("Saved inserted paragraph verification failed.")
        if table_hashes(verified) != before_tables:
            raise RuntimeError("An Appendix table changed during paragraph insertion.")
        if bookmark_fingerprint(verified) != before_bookmarks:
            raise RuntimeError("Bookmark fingerprint changed during paragraph insertion.")

        after_members = package_members(staged)
        changed = sorted(
            name
            for name in set(before_members) | set(after_members)
            if before_members.get(name) != after_members.get(name)
        )
        forbidden = [
            name
            for name in changed
            if name.startswith("word/") and name != "word/document.xml"
        ]
        if forbidden:
            raise RuntimeError(f"Protected Word members changed: {forbidden}")

        receipt: dict[str, object] = {
            "status": "dry-run" if dry_run else "applied",
            "sha256_before": before_hash,
            "sha256_after": sha256(staged),
            "changed_package_members": changed,
            "paragraphs_inserted": 1,
            "tables_unchanged": True,
            "bookmarks_preserved": True,
            "backup": None,
        }
        if dry_run:
            return receipt

        backup_dir = docx.parent / ".kila-backups"
        backup_dir.mkdir(parents=True, exist_ok=True)
        stamp = dt.datetime.now().strftime("%Y%m%dT%H%M%S%f")
        backup = backup_dir / f"Appendix.before-r1c1.{stamp}.docx"
        shutil.copy2(docx, backup)
        if sha256(backup) != before_hash:
            raise RuntimeError("Backup hash mismatch.")
        os.replace(staged, docx)
        receipt["backup"] = str(backup)
        return receipt
    finally:
        if staged.exists():
            staged.unlink()


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--docx", type=Path, default=Path("Rev/revision/Appendix.docx"))
    parser.add_argument("--apply", action="store_true")
    args = parser.parse_args()
    print(json.dumps(update(args.docx, dry_run=not args.apply), indent=2))


if __name__ == "__main__":
    main()
